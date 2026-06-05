"""
update_td_with_iflow: Updates an existing TD document with iFlow ZIP data.
- Copies Appendix data to main body tables
- Overrides iFlow-specific fields from the actual ZIP
- Adds iFlow design steps section
- Adds message mapping field table
- Adds SAP-themed diagram
ZERO AI — 100% deterministic.
"""

import io, re, zipfile
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from services.flowchart_builder import generate_flowchart
from services.iflow_parser import parse_iflow_zip as _parse_iflow_zip_full

# ─────────────────────────────────────────────────────────────────────────────
# PACKAGE NAME DERIVATION
# Convention:
#   Single country → {ISO2} - {VendorName}   e.g.  IT - ITALTRANS
#   Multi-country  → GLO - {VendorName}       e.g.  GLO - ITALTRANS
# ─────────────────────────────────────────────────────────────────────────────

# Country keyword → ISO-2 code
# Ordered: longer/more-specific names first to avoid short-name false matches
_COUNTRY_MAP: list[tuple[str, str]] = [
    # Europe
    ("ITALY",          "IT"), ("ITAL",          "IT"), ("ITALIAN",       "IT"),
    ("SPAIN",          "ES"), ("SPANISH",        "ES"), ("ESPANA",        "ES"),
    ("GERMANY",        "DE"), ("GERMAN",         "DE"), ("DEUTSCH",       "DE"),
    ("FRANCE",         "FR"), ("FRENCH",         "FR"), ("FRANCAIS",      "FR"),
    ("POLAND",         "PL"), ("POLISH",         "PL"), ("POLSKA",        "PL"),
    ("NETHERLANDS",    "NL"), ("DUTCH",          "NL"), ("HOLLAND",       "NL"),
    ("BELGIUM",        "BE"), ("BELGIAN",        "BE"),
    ("PORTUGAL",       "PT"), ("PORTUGUESE",     "PT"),
    ("SWEDEN",         "SE"), ("SWEDISH",        "SE"),
    ("NORWAY",         "NO"), ("NORWEGIAN",      "NO"),
    ("DENMARK",        "DK"), ("DANISH",         "DK"),
    ("FINLAND",        "FI"), ("FINNISH",        "FI"),
    ("AUSTRIA",        "AT"), ("AUSTRIAN",       "AT"),
    ("SWITZERLAND",    "CH"), ("SWISS",          "CH"),
    ("GREECE",         "GR"), ("GREEK",          "GR"),
    ("CZECHIA",        "CZ"), ("CZECH",          "CZ"),
    ("HUNGARY",        "HU"), ("HUNGARIAN",      "HU"),
    ("ROMANIA",        "RO"), ("ROMANIAN",       "RO"),
    ("BULGARIA",       "BG"), ("BULGARIAN",      "BG"),
    ("CROATIA",        "HR"), ("CROATIAN",       "HR"),
    ("SLOVAKIA",       "SK"), ("SLOVAK",         "SK"),
    ("SLOVENIA",       "SI"), ("SLOVENIAN",      "SI"),
    ("UKRAINE",        "UA"), ("UKRAINIAN",      "UA"),
    ("RUSSIA",         "RU"), ("RUSSIAN",        "RU"),
    ("TURKEY",         "TR"), ("TURKISH",        "TR"),
    ("UNITED KINGDOM", "GB"), ("UK",             "GB"), ("BRITAIN",       "GB"),
    # Americas
    ("UNITED STATES",  "US"), ("USA",            "US"), ("AMERICA",       "US"),
    ("CANADA",         "CA"), ("CANADIAN",       "CA"),
    ("MEXICO",         "MX"), ("MEXICAN",        "MX"),
    ("BRAZIL",         "BR"), ("BRAZILIAN",      "BR"),
    ("ARGENTINA",      "AR"), ("ARGENTINIAN",    "AR"),
    ("CHILE",          "CL"), ("CHILEAN",        "CL"),
    ("COLOMBIA",       "CO"), ("COLOMBIAN",      "CO"),
    # Asia-Pacific
    ("CHINA",          "CN"), ("CHINESE",        "CN"),
    ("JAPAN",          "JP"), ("JAPANESE",       "JP"),
    ("INDIA",          "IN"), ("INDIAN",         "IN"),
    ("KOREA",          "KR"), ("KOREAN",         "KR"),
    ("AUSTRALIA",      "AU"), ("AUSTRALIAN",     "AU"),
    ("SINGAPORE",      "SG"), ("SINGAPOREAN",    "SG"),
    ("MALAYSIA",       "MY"), ("MALAYSIAN",      "MY"),
    ("INDONESIA",      "ID"), ("INDONESIAN",     "ID"),
    ("THAILAND",       "TH"), ("THAI",           "TH"),
    ("VIETNAM",        "VN"), ("VIETNAMESE",     "VN"),
    # Middle East / Africa
    ("UAE",            "AE"), ("EMIRATES",       "AE"), ("DUBAI",         "AE"),
    ("SAUDI",          "SA"), ("ARABIA",         "SA"),
    ("EGYPT",          "EG"), ("EGYPTIAN",       "EG"),
    ("SOUTH AFRICA",   "ZA"),
    ("NIGERIA",        "NG"), ("NIGERIAN",       "NG"),
]

# SAP system names to EXCLUDE from vendor detection
_SAP_SYSTEMS = {
    "SAP", "S4HANA", "S4", "CPI", "AEM", "BTP", "ERP", "ECC", "PI", "PO",
    "INTEGRATION", "PROCESS", "CLOUD", "HANA", "FIORI", "ABAP", "BASIS",
}

# Prefixes to strip from vendor participant names to get clean vendor name
_VENDOR_STRIP_PREFIXES = ("3PL", "3PL ", "SYSTEM", "PARTNER", "EXTERNAL", "RECEIVER")


def _detect_countries(text: str) -> set[str]:
    """Find all country ISO-2 codes mentioned in the given text."""
    upper = text.upper()
    found = set()
    for keyword, iso in _COUNTRY_MAP:
        # Word-boundary match to avoid false positives (e.g. "IT" inside "WITH")
        if re.search(r'\b' + re.escape(keyword) + r'\b', upper):
            found.add(iso)
    return found


def _extract_vendor_name(facts: dict) -> str:
    """
    Extract the external vendor/partner name from the iFlow facts.
    Prefers the first non-SAP receiver participant name.
    Falls back to extracting from the iFlow Bundle-Name.
    """
    # Try participants (receivers that aren't SAP systems)
    for p in facts.get('participants', []):
        name = p.get('name', '').strip()
        ptype = p.get('type', '')
        if not name or ptype == 'IntegrationProcess':
            continue
        name_upper = name.upper().replace('_', '')
        if any(sap in name_upper for sap in _SAP_SYSTEMS):
            continue
        if name_upper in ('TRIGGERSALERTMAIL', 'TRIGGERALERTMAIL', 'MAIL', 'EMAIL'):
            continue
        # Strip common prefixes from vendor name
        display = name
        for prefix in _VENDOR_STRIP_PREFIXES:
            if display.upper().startswith(prefix):
                display = display[len(prefix):].strip()
        return display.upper() if display else name.upper()

    # Fallback: extract from iFlow Bundle-Name (e.g. "Send - Batch Status to 3PL ITALTRANS")
    iflow_name = facts.get('iflow_name', '')
    if iflow_name:
        # Take the last word(s) — usually the vendor name
        words = iflow_name.split()
        for i, word in enumerate(reversed(words)):
            w_upper = word.upper()
            if (len(w_upper) >= 3 and
                not any(sap in w_upper for sap in _SAP_SYSTEMS) and
                w_upper not in ('AND', 'TO', 'FROM', 'VIA', 'THE', 'A', 'AN')):
                return w_upper

    return ''


def derive_package_name(facts: dict, extra_text: str = '') -> str:
    """
    Derive the CPI package name following the convention:
      Single country  → {ISO2} - {VendorName}   e.g.  IT - ITALTRANS
      Multi-country   → GLO - {VendorName}       e.g.  GLO - ITALTRANS
      Unknown country → GLO - {VendorName}

    Sources searched (in priority order):
    1. iFlow Bundle-Name
    2. metainfo.prop description
    3. Participant names (receivers)
    4. Step names
    5. extra_text (e.g. TD appendix data — contains "Target system: Italy Italtrans 3PL")
    """
    # Collect all text to search for country keywords
    search_text = ' '.join(filter(None, [
        facts.get('iflow_name', ''),
        facts.get('description', ''),
        ' '.join(p.get('name', '') for p in facts.get('participants', [])),
        ' '.join(s.get('name', '') for s in facts.get('steps', [])),
        ' '.join(a.get('name', '') or a.get('url', '') for a in facts.get('adapters', [])),
        extra_text,
    ]))

    countries = _detect_countries(search_text)
    vendor    = _extract_vendor_name(facts)

    if not vendor:
        return 'GLO'

    if len(countries) == 0:
        prefix = 'GLO'
    elif len(countries) == 1:
        prefix = list(countries)[0]
    else:
        prefix = 'GLO'

    return f'{prefix} - {vendor}'

# ── Colour helpers ────────────────────────────────────────────────────────────
SAP_BLUE = RGBColor(0x00, 0x6D, 0xB3)
SAP_DARK = RGBColor(0x1A, 0x1A, 0x2E)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
AMBER    = RGBColor(0xCC, 0x77, 0x00)
MID_GREY = RGBColor(0x55, 0x60, 0x70)
GREEN    = RGBColor(0x1A, 0x8A, 0x5A)
RED_C    = RGBColor(0xC0, 0x2C, 0x2C)


def _shd(cell, hex_color):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), hex_color); s.set(qn("w:color"), "auto"); s.set(qn("w:val"), "clear")
    p.append(s)


def _set_cell(cell, text, bold=False, color=None):
    """Set cell text, preserve existing formatting where possible."""
    para = cell.paragraphs[0]
    for run in para.runs:
        run.text = ""
    if not para.runs:
        run = para.add_run()
    else:
        run = para.runs[0]
    run.text = str(text)
    run.font.size = Pt(9)
    run.font.bold = bold
    if color: run.font.color.rgb = color


def _norm(text):
    """Normalize label text for matching."""
    return re.sub(r'[^a-z0-9]', '', text.lower())[:30]


def _extract_iflow_facts(zip_bytes: bytes) -> dict:
    """Extract all key facts from the iFlow ZIP."""
    facts = {
        'iflow_name': '',
        'iflow_symbolic': '',
        'iflow_version': '',
        'description': '',
        'package': '',
        'mmap_name': '',
        'mmap_dst_fields': [],
        'mmap_src_fields': [],
        'mmap_xml': '',
        'steps': [],
        'adapters': [],
        'participants': [],
        'scripts': {},
        'parameters': [],
        'xsd_files': [],
    }

    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
        files = z.namelist()

        # MANIFEST.MF
        for fn in files:
            if 'MANIFEST.MF' in fn:
                mf = z.read(fn).decode('utf-8', 'replace')
                bn = re.search(r'^Bundle-Name:\s*(.+)', mf, re.MULTILINE)
                bs = re.search(r'^Bundle-SymbolicName:\s*(.+)', mf, re.MULTILINE)
                bv = re.search(r'^Bundle-Version:\s*(.+)', mf, re.MULTILINE)
                facts['iflow_name']     = _join_wrapped(bn.group(1).strip() if bn else '')
                facts['iflow_symbolic'] = _join_wrapped(bs.group(1).strip() if bs else '')
                facts['iflow_version']  = bv.group(1).strip() if bv else ''
                break

        # metainfo.prop
        for fn in files:
            if 'metainfo.prop' in fn:
                mp = z.read(fn).decode('utf-8', 'replace')
                for line in mp.splitlines():
                    if line.startswith('description='):
                        facts['description'] = line[12:].strip()

        # parameters.propdef
        for fn in files:
            if fn.endswith('.propdef'):
                raw = z.read(fn).decode('utf-8', 'replace')
                for m in re.finditer(r'<parameter>(.*?)</parameter>', raw, re.DOTALL):
                    blk = m.group(1)
                    nm  = re.search(r'<name>([^<]+)</name>', blk)
                    dsc = re.search(r'<description>([^<]+)</description>', blk)
                    req = re.search(r'<isRequired>([^<]+)</isRequired>', blk)
                    if nm:
                        facts['parameters'].append({
                            'name': nm.group(1).strip(),
                            'description': dsc.group(1).strip() if dsc else '',
                            'required': req.group(1).strip().lower() == 'true' if req else False,
                        })

        # Groovy scripts
        for fn in files:
            if fn.endswith('.groovy'):
                facts['scripts'][fn.split('/')[-1]] = z.read(fn).decode('utf-8', 'replace')

        # XSD files
        for fn in files:
            if fn.endswith('.xsd'):
                facts['xsd_files'].append(fn.split('/')[-1])

        # .mmap files
        mmap_files = [n for n in files if n.endswith('.mmap')]
        if mmap_files:
            mf = mmap_files[0]
            facts['mmap_name'] = mf.split('/')[-1].replace('.mmap', '')
            mmap_xml = z.read(mf).decode('utf-8', 'replace')
            facts['mmap_xml'] = mmap_xml
            dst = re.findall(r'path="([^"]+)" type="Dst"', mmap_xml)
            src = re.findall(r'path="([^"]+)" type="Src"', mmap_xml)
            facts['mmap_dst_fields'] = [d.split('/')[-1] for d in dst]
            facts['mmap_src_fields'] = [s.split('/')[-1] for s in src]

        # iflw XML
        iflw_files = [n for n in files if n.endswith('.iflw')]
        if iflw_files:
            xml = z.read(iflw_files[0]).decode('utf-8', 'replace')
            
            # Participants
            for m in re.finditer(r'<bpmn2:participant\b([^>]*)>', xml):
                attrs = m.group(1)
                nm  = re.search(r'\bname="([^"]+)"', attrs)
                tp  = re.search(r'\bifl:type="([^"]+)"', attrs)
                if nm:
                    facts['participants'].append({'name': nm.group(1), 'type': tp.group(1) if tp else ''})

            # Steps (all element types)
            step_patterns = [
                ('startEvent',        r'<bpmn2:startEvent\b([^>]*)>(.*?)</bpmn2:startEvent>'),
                ('callActivity',      r'<bpmn2:callActivity\b([^>]*)>(.*?)</bpmn2:callActivity>'),
                ('serviceTask',       r'<bpmn2:serviceTask\b([^>]*)>(.*?)</bpmn2:serviceTask>'),
                ('exclusiveGateway',  r'<bpmn2:exclusiveGateway\b([^>]*)>(.*?)</bpmn2:exclusiveGateway>'),
                ('subProcess',        r'<bpmn2:subProcess\b([^>]*)>(.*?)</bpmn2:subProcess>'),
                ('endEvent',          r'<bpmn2:endEvent\b([^>]*)>(.*?)</bpmn2:endEvent>'),
            ]
            for etype, pat in step_patterns:
                for m in re.finditer(pat, xml, re.DOTALL):
                    attrs = m.group(1); body = m.group(2)
                    nm  = re.search(r'\bname="([^"]+)"', attrs)
                    at  = re.search(r'<key>activityType</key><value>([^<]+)</value>', body)
                    sc  = re.search(r'<key>script</key><value>([^<]+)</value>', body)
                    cmd = re.search(r'<key>cmdVariantUri</key><value>([^<]+)</value>', body)
                    facts['steps'].append({
                        'name': nm.group(1) if nm else '',
                        'element_type': etype,
                        'activity_type': at.group(1).strip() if at else '',
                        'script_file': sc.group(1).strip() if sc else '',
                        'cmd_uri': cmd.group(1).strip() if cmd else '',
                    })

            # Adapters from messageFlows
            for m in re.finditer(r'<bpmn2:messageFlow\b([^>]*)>(.*?)</bpmn2:messageFlow>', xml, re.DOTALL):
                attrs = m.group(1); body = m.group(2)
                nm   = re.search(r'\bname="([^"]+)"', attrs)
                src_ref = re.search(r'\bsourceRef="([^"]+)"', attrs)
                tgt_ref = re.search(r'\btargetRef="([^"]+)"', attrs)
                ct  = re.search(r'<key>ComponentType</key><value>([^<]+)</value>', body)
                dir_ = re.search(r'<key>direction</key><value>([^<]+)</value>', body)
                url = re.search(r'<key>(?:address|httpAddressWithoutQuery|urlPath)</key><value>([^<]+)</value>', body)
                cred= re.search(r'<key>(?:credentialName|credential_name|alias)</key><value>([^<]+)</value>', body)
                facts['adapters'].append({
                    'name': nm.group(1) if nm else '',
                    'component': ct.group(1).strip() if ct else '',
                    'direction': dir_.group(1).strip() if dir_ else '',
                    'url': url.group(1).strip() if url else '',
                    'credential': cred.group(1).strip() if cred else '',
                    'source_ref': src_ref.group(1) if src_ref else '',
                    'target_ref': tgt_ref.group(1) if tgt_ref else '',
                })

    return facts


def _join_wrapped(text):
    """Join MANIFEST continuation lines (lines starting with space)."""
    return re.sub(r'\s+', ' ', text.replace('\n ', '')).strip()


def _extract_appendix_data(doc: Document) -> dict:
    """
    Extract data from appendix tables (last ~14 tables in the document).
    Returns dict: {normalized_label: value}
    """
    data = {}
    n = len(doc.tables)
    # Appendix tables: widen to last ~20 to capture business-info tables
    # that sometimes sit just before the technical appendix block
    appendix_start = max(0, n - 20)
    
    for table in doc.tables[appendix_start:]:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Deduplicate merged
            seen = []
            for c in cells:
                if c not in seen: seen.append(c)
            if len(seen) >= 2 and seen[0] and seen[1]:
                label = seen[0]
                value = seen[1]
                if label and value and value not in ('', 'N/A', '[TBD]', '[TBD Date]', '[TBD Name]'):
                    data[_norm(label)] = value
    return data


def _check_middleware_checkboxes(element_or_doc, solutions: set):
    """
    Tick the W14 SDT checkbox controls that precede matching solution labels.
    element_or_doc: doc (whole document) OR a table._tbl element (single table only)
    solutions: e.g. {'CPI', 'Event Mesh', 'SAP System', 'HTTPS'}
    """
    root = element_or_doc.element.body if hasattr(element_or_doc, 'element') else element_or_doc
    _W14 = 'http://schemas.microsoft.com/office/word/2010/wordml'

    def _w14(tag):
        return f'{{{_W14}}}{tag}'

    for sdt in root.iter(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        cb = sdtPr.find(_w14('checkbox'))
        if cb is None:
            continue  # not a checkbox SDT

        # Read the label in the sibling run(s) that immediately follow this SDT
        parent = sdt.getparent()
        if parent is None:
            continue
        siblings = list(parent)
        try:
            idx = siblings.index(sdt)
        except ValueError:
            continue

        label = ''
        for sib in siblings[idx + 1: idx + 4]:
            for t in sib.iter(qn('w:t')):
                label += (t.text or '')
            if label.strip():
                break

        # Does this checkbox's label match any solution we want to tick?
        label_u = label.strip().upper()
        if not any(s.upper() in label_u for s in solutions):
            continue

        # ── Mark checked ──────────────────────────────────────────────────────
        checked_elem = cb.find(_w14('checked'))
        if checked_elem is not None:
            checked_elem.set(_w14('val'), '1')

        # Update the visible character inside sdtContent  ☐ → ☑
        sdt_content = sdt.find(qn('w:sdtContent'))
        if sdt_content is not None:
            for t_elem in sdt_content.iter(qn('w:t')):
                if t_elem.text in ('☐', '□', '◻', ''):
                    t_elem.text = '☑'  # ☑ Ballot Box With Check Mark


# Labels where we ALWAYS replace the current value (even if not empty)
# because the TD may have wrong placeholder names from the FD
_ALWAYS_REPLACE_LABELS = {
    'iflow',           # IFlow name
    'artifact',        # Artifact (CPI)
    'package',         # Package / Folder name
    'folder',          # Folder name (PO)
    'workingname',     # Working Name of Interface → iFlow name
    'mappingname',     # Mapping name row
    'mmapping',        # Mapping name
    'description',     # iFlow description
    'mode',            # Asynchronous/Synchronous
    'bundlename',      # Bundle name
    'artifactname',    # Artifact name
}

def _should_always_replace(norm_label: str) -> bool:
    """Return True if this label's value should always be replaced from iFlow ZIP."""
    return any(k in norm_label for k in _ALWAYS_REPLACE_LABELS)


def _fill_table_from_appendix(table, appendix_data: dict, overrides: dict = None):
    """
    For each row in a main-body table:
    - If label is in ALWAYS_REPLACE list → replace regardless of current value
    - If cell is empty or has a placeholder → fill from appendix / overrides
    overrides: {normalized_label: value} that takes priority over appendix.
    """
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        label_text = cells[0].text.strip()
        if not label_text:
            continue
        norm_label = _norm(label_text)
        
        # Find the VALUE cell (skip merged cells, find first non-label empty-ish cell)
        value_cell = None
        for c in cells[1:]:
            if c != cells[0]:  # skip merged
                value_cell = c
                break
        if not value_cell:
            continue

        current_val = value_cell.text.strip()
        norm_label  = _norm(label_text)

        # Always replace for iFlow-critical fields (iFlow name, package, mmap name, description)
        force_replace = _should_always_replace(norm_label) and bool(overrides)

        # For other fields: only fill if empty or placeholder
        is_empty = (not current_val or
                    current_val.startswith('[TBD') or
                    current_val.startswith('Example:') or
                    current_val.startswith('Type TBU') or
                    current_val.startswith('Make a summary') or
                    current_val.startswith('Place a BPMN') or
                    current_val.startswith('TBF') or
                    current_val.startswith('upon receiving'))

        if not is_empty and not force_replace:
            continue

        # Check overrides first (iFlow ZIP data), then appendix
        new_val = None
        if overrides:
            for ok, ov in overrides.items():
                if ok in norm_label or norm_label in ok:
                    new_val = ov
                    break
        if not new_val:
            for ak, av in appendix_data.items():
                if ak in norm_label or norm_label in ak:
                    new_val = av
                    break

        if new_val:
            _set_cell(value_cell, new_val)


def _add_para(doc, text, bold=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    if color: r.font.color.rgb = color
    return p


def _add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        r = c.paragraphs[0].runs[0]; r.font.bold = True
        r.font.size = Pt(9); r.font.color.rgb = WHITE
        _shd(c, '1A1A2E')
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val) if val else ''
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0: _shd(c, 'F2F4F7')
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows: row.cells[j].width = Cm(w)
    doc.add_paragraph()


def _add_code(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line or ' ')
        r.font.name = 'Courier New'; r.font.size = Pt(8); r.font.color.rgb = SAP_DARK


STEP_LABELS = {
    'Script':                      'Groovy Script',
    'GroovyScript':                'Groovy Script',
    'Enricher':                    'Content Modifier',
    'ExternalCall':                'Request Reply',
    'Mapping':                     'Message Mapping',
    'ExclusiveGateway':            'Router (CBR)',
    'Splitter':                    'Splitter',
    'DBstorage':                   'DataStore',
    'StartTimerEvent':             'Timer Start',
    'ErrorEventSubProcessTemplate':'Exception Subprocess',
    'Send':                        'Send Step',
    'ProcessCallElement':          'Process Call',
}


def _replace_text_in_cell(cell, old_text: str, new_text: str):
    """Replace text in all runs of all paragraphs in a cell."""
    for para in cell.paragraphs:
        full = ''.join(r.text for r in para.runs)
        if old_text in full:
            new_full = full.replace(old_text, new_text)
            if para.runs:
                para.runs[0].text = new_full
                for r in para.runs[1:]:
                    r.text = ''
            else:
                para.add_run(new_full)


def _replace_all_occurrences(doc: Document, replacements: list[tuple[str, str]]):
    """
    Replace ALL occurrences of old→new text across EVERY run in EVERY paragraph
    in EVERY table cell and regular paragraph in the document.
    replacements: list of (old_text, new_text)
    """
    def _fix_para(para):
        for old, new in replacements:
            full = ''.join(r.text for r in para.runs)
            if old in full:
                new_full = full.replace(old, new)
                if para.runs:
                    para.runs[0].text = new_full
                    for r in para.runs[1:]:
                        r.text = ''
                else:
                    para.add_run(new_full)

    # Regular paragraphs
    for para in doc.paragraphs:
        _fix_para(para)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fix_para(para)


def _infer_step_purpose(step: dict, facts: dict) -> tuple[str, str, list[str]]:
    """
    Returns (purpose_sentence, developer_note, key_config_list) for a step.
    Infers purpose from naming conventions, step type, and script content.
    """
    name = step.get('name', '')
    act  = step.get('activity_type', '')
    elem = step.get('element_type', '')
    props = step.get('properties', {})
    script_code = step.get('script_code', '')
    name_upper = name.upper().replace('_', ' ').replace('-', ' ')

    # ── Infer activity type from naming convention if not set ─────────────────
    # SAP CPI naming: GS_=Groovy, CM_=Content Modifier, MM_=MessageMapping,
    # RR_=RequestReply, MR_=MessageRouter
    if not act:
        nu = name.upper()
        if nu.startswith('GS_') or nu.startswith('GS ') or 'GROOVY' in nu:
            act = 'Script'
        elif nu.startswith('CM_') or nu.startswith('CM '):
            act = 'Enricher'
        elif nu.startswith('MM_') or nu.startswith('MM ') or 'MAPPING' in nu and 'MESSAGE' in nu:
            act = 'Mapping'
        elif nu.startswith('RR_') or nu.startswith('RR '):
            act = 'ExternalCall'
        elif nu.startswith('MR_') or 'ROUTER' in nu or 'GATEWAY' in nu:
            act = 'ExclusiveGateway'
        elif 'JSON' in nu and 'XML' in nu:
            act = 'JsonToXmlConverter'
        elif 'XML' in nu and 'JSON' in nu:
            act = 'XmlToJsonConverter'
        elif 'REMOVE XML' in nu or 'REMOVEXML' in nu:
            act = 'Enricher'

    purpose = ''
    dev_note = ''
    key_config = []

    # ── Start/End events ──────────────────────────────────────────────────────
    if elem == 'startEvent':
        if 'Error' in name:
            purpose = 'Entry point of the Exception Subprocess. Activated whenever an uncaught exception occurs in the main integration flow.'
            dev_note = 'SAP CPI automatically routes to this on any runtime exception.'
        else:
            _senders = [a for a in facts.get('adapters', []) if a['direction'] == 'Sender']
            src  = (_senders[0].get('source_name') or _senders[0].get('system','')
                    or 'source system') if _senders else 'source system'
            comp = _senders[0]['component'] if _senders else 'HTTPS'
            purpose = (f'Integration entry point. Receives the inbound message from {src} '
                       f'via the {comp} adapter. The message arrives as a trigger from the sender system '
                       f'and initiates the integration processing chain.')
            dev_note = f'In the Designer: drag a Message Start Event. Connect the sender participant via a {comp} adapter messageFlow.'
    elif elem == 'endEvent':
        purpose = 'Marks the successful completion of the integration flow. The message has been processed and delivered.'
        dev_note = 'Drag a Message End Event after the last processing step.'

    # ── Groovy Scripts ────────────────────────────────────────────────────────
    elif act in ('Script', 'GroovyScript'):
        fn = step.get('script_fn', 'processData')
        sf = step.get('script_file', '')

        # Infer from name keywords
        if 'LOGBEFORE' in name_upper or 'LOG BEFORE' in name_upper:
            purpose = ('Captures and logs the incoming message payload BEFORE any transformation. '
                       'Creates a Message Processing Log (MPL) attachment for audit trail and troubleshooting. '
                       'Only active when ENABLE_PAYLOAD_LOGGING parameter is set to TRUE.')
            dev_note = 'Set ENABLE_PAYLOAD_LOGGING=TRUE in test environments. Keep FALSE in production for performance.'
            key_config = ['Parameter: ENABLE_PAYLOAD_LOGGING (TRUE/FALSE)', f'Script: {sf}', f'Function: {fn}']

        elif 'LOGAFTER' in name_upper or 'LOG AFTER' in name_upper:
            purpose = ('Captures and logs the transformed message payload AFTER mapping/conversion. '
                       'Allows comparison of before/after state in the MPL for debugging. '
                       'Controlled by the same ENABLE_PAYLOAD_LOGGING parameter.')
            key_config = ['Parameter: ENABLE_PAYLOAD_LOGGING (TRUE/FALSE)', f'Script: {sf}']

        elif 'LOGEXCEPTION' in name_upper or 'LOG EXCEPTION' in name_upper or 'LOGPAYLOAD' in name_upper and 'Exception' in name:
            purpose = ('Logs the failed payload and full exception stack trace into the MPL when an error occurs. '
                       'Provides developers and operations with the exact payload that caused the failure '
                       'for root cause analysis and manual retry decisions.')
            key_config = [f'Script: {sf}', 'Logs: ErrorMessage, ErrorClass, FailedPayload attachment']

        elif 'SETEXCEPTION' in name_upper or 'SET EXCEPTION' in name_upper:
            purpose = ('Extracts the exception details (error class, error message) from the caught exception '
                       'and sets them as message properties. These properties are then used by the email '
                       'notification step to include meaningful error information in the alert.')
            key_config = ['Sets: ErrorMessage, ErrorClass properties', f'Script: {sf}']

        elif 'SETPROP' in name_upper or 'SET PROP' in name_upper:
            purpose = 'Sets required message properties for downstream processing.'
            key_config = [f'Script: {sf}', f'Function: {fn}']

        else:
            # Read first meaningful comment from script code
            desc_from_code = ''
            if script_code:
                for line in script_code.split('\n')[:20]:
                    line = line.strip()
                    if line.startswith('//') or line.startswith('*') or line.startswith('/**'):
                        clean = line.lstrip('/* ').rstrip(' */')
                        if len(clean) > 20 and not clean.startswith('@'):
                            desc_from_code = clean
                            break

            purpose = desc_from_code or f'Groovy script step: {name}. Performs custom processing logic.'
            dev_note = f'Script file: {sf}, function: {fn}(Message msg)'
            key_config = [f'Script: {sf}', f'Function: {fn}']

    # ── Content Modifier (Enricher) ───────────────────────────────────────────
    elif act == 'Enricher':
        header_tbl = props.get('headerTable', '')
        prop_tbl   = props.get('propertyTable', '')

        # Extract what's being set
        headers_set = re.findall(r"<cell id='Name'>([^<]+)</cell>", header_tbl)
        props_set   = re.findall(r"<cell id='Name'>([^<]+)</cell>", prop_tbl)

        if 'DELETE' in name_upper or 'REMOVE HEADER' in name_upper:
            purpose = ('Removes all HTTP headers inherited from the incoming message. '
                       'This is a best practice in SAP CPI to prevent unwanted headers '
                       '(e.g., Authorization, Content-Type from AEM) from being passed to downstream receivers. '
                       'Ensures the adapter configurations control the outbound headers, not the inbound ones.')
            dev_note = 'In Content Modifier: Message Header tab → set Action=Delete for all headers, or use a wildcard.'

        elif 'SETATTR' in name_upper or 'SET ATTR' in name_upper or 'SETATTRIB' in name_upper:
            if props_set:
                purpose = (f'Sets message properties required for processing: {", ".join(props_set[:6])}. '
                           'These properties are used by downstream steps for routing, adapter configuration, '
                           'and monitoring. Must be set before the routing gateway.')
            else:
                purpose = 'Sets message attributes (properties/headers) required for downstream processing and adapter configuration.'
            if props_set:
                key_config = [f'Sets properties: {", ".join(props_set[:5])}']
            if headers_set:
                key_config.append(f'Sets headers: {", ".join(headers_set[:5])}')
            dev_note = 'Content Modifier: Exchange Properties tab. Use expressions like ${header.x} or ${property.y}.'

        elif 'CUSTOMSTATUS' in name_upper or 'UPDATE' in name_upper and 'STATUS' in name_upper:
            purpose = ('Updates the SAP_MessageProcessingLogCustomStatus property, which is used by '
                       'SAP Cloud ALM and Message Monitor for operational tracking. '
                       'Setting a meaningful status (e.g., COMPLETED with business key) allows the '
                       'operations team to search and filter messages by business context.')
            status_val = re.search(r"Value'>(COMPLETED[^<]*)</cell>", prop_tbl)
            if status_val:
                key_config = [f'Status value: {status_val.group(1)[:60]}']
            dev_note = 'Required for Cloud ALM integration. Property: SAP_MessageProcessingLogCustomStatus'

        elif 'SETBODY' in name_upper or 'SET BODY' in name_upper:
            purpose = 'Sets or transforms the message body content for the outbound payload.'
            dev_note = 'Content Modifier: Body tab. Can use expressions or fixed content.'

        elif headers_set or props_set:
            all_set = (headers_set + props_set)[:6]
            purpose = f'Content Modifier that sets the following: {", ".join(all_set)}.'
            key_config = [f'Headers set: {", ".join(headers_set[:3])}'] if headers_set else []
            if props_set: key_config.append(f'Properties set: {", ".join(props_set[:3])}')

        else:
            purpose = f'Content Modifier step: {name}. Configures message headers and/or properties.'

    # ── Message Mapping ───────────────────────────────────────────────────────
    elif act == 'Mapping':
        mmap = step.get('script_file', '') or facts.get('mmap_name', '')
        xsds = facts.get('xsd_files', [])
        src_xsd = next((x for x in xsds if not x.startswith('Z_')), xsds[0] if xsds else '')
        tgt_xsd = next((x for x in xsds if x.startswith('Z_')), xsds[-1] if xsds else '')
        purpose = (f'Performs the structural IDoc-to-IDoc message mapping using the '
                   f'SAP CPI Graphical Message Mapping artifact "{mmap}". '
                   f'Transforms the standard SAP format ({src_xsd}) into the customer-specific '
                   f'format ({tgt_xsd}) required by the receiver. '
                   f'Field mappings include {len(facts.get("mmap_dst_fields", []))} target fields '
                   f'from {len(facts.get("mmap_src_fields", []))} source fields.')
        dev_note = (f'Message Mapping artifact "{mmap}" must be deployed separately in the same CPI package '
                    f'before this iFlow can be deployed. Upload the .mmap ZIP to the package first.')
        key_config = [f'Mapping: {mmap}', f'Source XSD: {src_xsd}', f'Target XSD: {tgt_xsd}']

    # ── JSON/XML Converters ───────────────────────────────────────────────────
    elif 'JsonToXml' in act or 'JSON' in name_upper and 'XML' in name_upper:
        purpose = ('Converts the incoming JSON payload (as sent by SAP Advanced Event Mesh) '
                   'into XML format. SAP CPI IDoc processing and message mapping require XML format. '
                   'AEM events are typically published in JSON; this step bridges the format gap.')
        dev_note = 'Palette: Call → JSON to XML Converter. No configuration usually needed.'

    elif 'XmlToJson' in act:
        purpose = 'Converts XML payload to JSON format for JSON-based receiver systems.'
        dev_note = 'Palette: Call → XML to JSON Converter.'

    # ── Router / Gateway ──────────────────────────────────────────────────────
    elif elem == 'exclusiveGateway':
        seqs = [f for f in facts.get('sequence_flows', [])
                if f.get('source') == step.get('id', '') and f.get('condition')]
        purpose = (f'Content-Based Router that evaluates conditions on message properties or headers '
                   f'to route the message to the correct downstream target. '
                   f'{"Conditions: " + "; ".join(f["condition"] for f in seqs[:3]) if seqs else "Routing conditions are configured on the outgoing sequence flows."}')
        dev_note = ('Palette: Routing → Router. Add outgoing routes. '
                    'Set condition type to Non-XML and write property expressions like ${property.receiverNumber} = "1".')
        if seqs:
            key_config = [f'Route "{f["name"]}": {f["condition"]}'
                          for f in seqs[:5] if f.get('name') and f.get('condition')]

    # ── Request Reply (Service Task) ──────────────────────────────────────────
    elif elem == 'serviceTask' or act == 'ExternalCall':
        adapters_for_step = [a for a in facts.get('adapters', [])
                              if a.get('source_name') or a.get('source_ref','') == step.get('id', '')
                              or a.get('name', '').replace(' ','') in name.replace(' ','')]
        if adapters_for_step:
            a = adapters_for_step[0]
            comp = a.get('component', 'unknown')
            tgt  = a.get('target_ref', '')
            purpose = (f'Request Reply step that delivers the message to the receiver system "{tgt}" '
                       f'using the {comp} adapter. This is the outbound delivery step where the '
                       f'transformed payload is sent to the target. '
                       f'{"Credential alias: " + a["credential"] if a.get("credential") else ""}')
            dev_note = f'Palette: Call → Request Reply. Use a bpmn2:serviceTask (NOT callActivity). Connect to {tgt} participant via {comp} adapter messageFlow.'
            cred = a.get('credential', '')
            url  = a.get('url', '')
            if cred: key_config.append(f'Credential alias: {cred}')
            if url:  key_config.append(f'Endpoint: {url[:60]}')
            key_config.append(f'Adapter: {comp}')
        elif 'MAIL' in name_upper or 'ALERT' in name_upper or 'EMAIL' in name_upper:
            purpose = ('Sends an email notification to the operations team when the integration fails. '
                       'Uses ProcessDirect to call a dedicated email-sending iFlow, keeping the error '
                       'notification logic separate from the main integration.')
            dev_note = 'ProcessDirect adapter: set address to match the consumer iFlow sender address.'
            key_config = ['Adapter: ProcessDirect', 'Calls: Email notification sub-iFlow']
        else:
            purpose = f'Outbound call step: {name}. Sends the message to the receiver.'
            key_config = [act]

    elif act == 'Send':
        adapters_for_step = [a for a in facts.get('adapters', [])
                              if a.get('source_name') or a.get('source_ref','') == step.get('id', '')]
        if adapters_for_step:
            a = adapters_for_step[0]
            purpose = (f'One-way send step that delivers the file/message to the receiver using {a.get("component","SFTP")} adapter. '
                       f'Unlike Request Reply, this does not wait for a response.')
        else:
            purpose = f'Outbound send step: {name}.'

    # ── Exception Subprocess ──────────────────────────────────────────────────
    elif act == 'ErrorEventSubProcessTemplate' or 'subprocess' in elem.lower():
        purpose = ('Exception Subprocess that catches ALL unhandled runtime exceptions from the main flow. '
                   'It runs automatically when any step in the main flow throws an exception. '
                   'Best practice in SAP CPI: always include an exception subprocess for proper '
                   'error logging, alerting, and MPL status management.')
        dev_note = ('Palette: Call → Exception Subprocess. Place at the bottom of the canvas. '
                    'Do NOT add triggeredByEvent=true attribute.')

    # ── Remove XML declaration ────────────────────────────────────────────────
    elif 'REMOVEXML' in name_upper or 'REMOVE XML' in name_upper or 'XML DEFINIT' in name_upper:
        purpose = ('Removes the XML declaration header (<?xml version="1.0" encoding="UTF-8"?>) '
                   'from the message payload. Some AS2 or B2B receivers cannot process messages '
                   'with an XML declaration prefix and expect raw content. '
                   'This step ensures compatibility with strict receiver requirements.')
        dev_note = 'Typically implemented as a Content Modifier with a Groovy expression on the Body tab.'

    # ── Fallback ──────────────────────────────────────────────────────────────
    if not purpose:
        label = STEP_LABELS.get(act, elem)
        purpose = f'{label} step "{name}". Performs processing as part of the integration flow.'

    return purpose, dev_note, key_config


def _build_narrative_steps(doc: Document, facts: dict, senders: list, receivers: list):
    """
    Build the main step-by-step narrative section.
    Each step gets a paragraph describing WHAT it does, WHY, and key config for development.
    """
    # ── Section heading ────────────────────────────────────────────────────────
    h3 = doc.add_heading('Step-by-Step Processing Guide', level=2)
    for r in h3.runs: r.font.color.rgb = SAP_BLUE

    _add_para(doc,
        'This section documents every step in the iFlow in execution order. '
        'A developer should follow these steps in SAP Integration Suite Designer '
        'to implement or recreate the iFlow from scratch.',
        size=10, color=MID_GREY)
    doc.add_paragraph()

    # ── Sender adapter narrative (before Step 1) ──────────────────────────────
    if senders:
        a = senders[0]
        h_send = doc.add_heading('Sender Adapter Configuration', level=3)
        for r in h_send.runs: r.font.color.rgb = RGBColor(0x00, 0x6B, 0x9F)
        _add_para(doc,
            f'The iFlow is triggered by the sender system "{a.get("source_name") or a.get("source_ref","")}" using the '
            f'{a["component"]} adapter. This defines HOW the message enters CPI.',
            size=10, bold=False)
        if a.get('url'):
            _add_para(doc, f'Endpoint path: {a["url"]}', size=9, color=MID_GREY, indent=0.5)
        if a.get('credential'):
            _add_para(doc, f'Authentication: credential alias "{a["credential"]}"', size=9, color=MID_GREY, indent=0.5)
        doc.add_paragraph()

    # ── Separate main flow from error handling ─────────────────────────────────
    main_steps = [s for s in facts['steps']
                  if s['element_type'] not in ('endEvent',)
                  and not ('Error' in s.get('name','') and s['element_type'] == 'startEvent')
                  and s.get('name')]

    error_steps = [s for s in facts['steps']
                   if (('Error' in s.get('name','') or 'Exception' in s.get('name',''))
                       and s['element_type'] not in ('endEvent',)
                       and s.get('name'))
                   or s.get('activity_type') == 'ErrorEventSubProcessTemplate']

    # Remove error steps from main using NAME matching (IDs may be empty)
    error_names = {s.get('name','').strip() for s in error_steps if s.get('name','').strip()}
    main_steps = [s for s in main_steps if s.get('name','').strip() not in error_names]

    # ── Main flow steps ────────────────────────────────────────────────────────
    h_main = doc.add_heading('Main Processing Flow', level=3)
    for r in h_main.runs: r.font.color.rgb = GREEN

    step_num = 1
    for s in main_steps:
        elem  = s['element_type']
        act   = s.get('activity_type', '')
        name  = s['name']
        label = STEP_LABELS.get(act, elem.replace('Event','').title())

        purpose, dev_note, key_config = _infer_step_purpose(s, facts)

        # Step heading
        color = ('1A8A5A' if 'Event' in elem
                 else 'CC7700' if elem == 'exclusiveGateway'
                 else '1A1A2E')
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        c0 = t.rows[0].cells[0]; c0.text = f'Step {step_num}'; c0.width = Cm(2.5)
        c1 = t.rows[0].cells[1]; c1.text = f'{name}  [{label}]'; c1.width = Cm(15)
        for ci, col_hex in [(c0, color),(c1, color)]:
            _shd(ci, col_hex)
            ci.paragraphs[0].runs[0].font.bold = True
            ci.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            ci.paragraphs[0].runs[0].font.size = Pt(10)
        step_num += 1

        # Purpose
        _add_para(doc, purpose, size=10, indent=0.3)

        # Key config
        if key_config:
            for cfg in key_config:
                _add_para(doc, f'  - {cfg}', size=9, color=MID_GREY)

        # Developer note
        if dev_note:
            _add_para(doc, f'Developer note: {dev_note}', size=9, color=AMBER, indent=0.3)

        doc.add_paragraph()

    # ── Receiver adapter narrative ─────────────────────────────────────────────
    real_receivers = [a for a in receivers if a.get('component') not in ('ProcessDirect',)]
    if real_receivers:
        h_recv = doc.add_heading('Receiver Adapter Configuration', level=3)
        for r in h_recv.runs: r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        for a in real_receivers:
            _add_para(doc,
                f'Delivers the transformed message to "{a.get("target_name") or a.get("target_ref","")}" via the {a["component"]} adapter.',
                size=10, bold=True)
            if a.get('url'):
                _add_para(doc, f'  Endpoint: {a["url"]}', size=9, color=MID_GREY)
            if a.get('credential'):
                _add_para(doc, f'  Credential alias: "{a["credential"]}" (create in CPI Security Material before deploying)',
                          size=9, color=AMBER)
            doc.add_paragraph()

    # ── Exception handling ─────────────────────────────────────────────────────
    if error_steps:
        h_err = doc.add_heading('Exception Handling Flow', level=3)
        for r in h_err.runs: r.font.color.rgb = RED_C

        _add_para(doc,
            'The following steps run ONLY when an exception occurs in the main flow. '
            'They are contained inside the Exception Subprocess at the bottom of the iFlow canvas.',
            size=10, color=MID_GREY)
        doc.add_paragraph()

        for s in error_steps:
            if s.get('activity_type') == 'ErrorEventSubProcessTemplate':
                continue  # container, skip
            elem  = s['element_type']
            act   = s.get('activity_type', '')
            name  = s.get('name','')
            if not name: continue
            label = STEP_LABELS.get(act, elem.replace('Event','').title())
            purpose, dev_note, key_config = _infer_step_purpose(s, facts)

            t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
            c0 = t.rows[0].cells[0]; c0.text = f'EX'; c0.width = Cm(1.5)
            c1 = t.rows[0].cells[1]; c1.text = f'{name}  [{label}]'; c1.width = Cm(16)
            for ci in (c0, c1):
                _shd(ci, 'C02C2C')
                ci.paragraphs[0].runs[0].font.bold = True
                ci.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                ci.paragraphs[0].runs[0].font.size = Pt(10)

            _add_para(doc, purpose, size=10, indent=0.3)
            if key_config:
                for cfg in key_config:
                    _add_para(doc, f'  - {cfg}', size=9, color=MID_GREY)
            if dev_note:
                _add_para(doc, f'Developer note: {dev_note}', size=9, color=AMBER, indent=0.3)
            doc.add_paragraph()

    # ── Externalized parameters (condensed, only required ones) ───────────────
    params = facts.get('parameters', [])
    required_params = [p for p in params if p.get('required')]
    optional_params = [p for p in params if not p.get('required')]

    if params:
        h_params = doc.add_heading('Configuration Parameters', level=3)
        for r in h_params.runs: r.font.color.rgb = SAP_BLUE
        _add_para(doc,
            'The following parameters must be configured in CPI after import: '
            'iFlow → Configure → Parameters tab.',
            size=10, color=MID_GREY)

        if required_params:
            _add_para(doc, 'Required (must fill before deploying):', bold=True, size=10)
            _add_table(doc,
                ['Parameter Name', 'Description'],
                [(p['name'], p['description'][:80] if p['description'] else 'Set before deploying')
                 for p in required_params[:20]],
                col_widths=[6, 11.5])

        if optional_params:
            _add_para(doc, f'Optional ({len(optional_params)} parameters — set as needed):', bold=False, size=9, color=MID_GREY)
            _add_table(doc,
                ['Parameter Name', 'Description'],
                [(p['name'], p['description'][:80] if p['description'] else '')
                 for p in optional_params[:15]],
                col_widths=[6, 11.5])


def _replace_artifact_names_everywhere(doc: Document, facts: dict, appendix_data: dict):
    """
    Scan every cell in every table in the document.
    Replace wrong/placeholder artifact names with correct names from iFlow ZIP.

    Targets:
    - Any cell whose label-row says 'IFlow', 'Artifact', 'Package', 'Folder name',
      'Name:' (in mapping context), 'Description'
    - The value cell in that row gets the correct name
    """
    iflow_name  = facts.get('iflow_name', '')
    mmap_name   = facts.get('mmap_name', '')
    pkg_name    = derive_package_name(facts)  # e.g. "IT - ITALTRANS"
    description = facts.get('description', '')

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells:
                continue

            # Get first cell as the label, subsequent non-merged cells as values
            label = cells[0].text.strip()
            if not label:
                continue
            nf = _norm(label)

            # Collect unique value cells
            value_cells = []
            seen_texts = {cells[0].text}
            for c in cells[1:]:
                if c.text not in seen_texts:
                    seen_texts.add(c.text)
                    value_cells.append(c)

            if not value_cells:
                continue
            val_cell = value_cells[0]
            curr = val_cell.text.strip()

            # ── iFlow name rows ───────────────────────────────────────────────
            if iflow_name and ('iflow' in nf or ('artifact' in nf and 'cpi' in nf)):
                if curr and curr != iflow_name:
                    _set_cell(val_cell, iflow_name)

            # ── Package / Folder name rows ────────────────────────────────────
            elif pkg_name and ('folder' in nf or 'package' in nf):
                if curr and curr != pkg_name:
                    _set_cell(val_cell, pkg_name)

            # ── Description rows (iFlow description) ─────────────────────────
            elif description and nf == _norm('Description:'):
                if not curr or curr.startswith('[') or len(curr) < 20:
                    _set_cell(val_cell, description)

            # ── Mapping name rows ─────────────────────────────────────────────
            # Replace if cell looks like an old wrong mapping/artifact name
            elif mmap_name:
                is_name_row = 'name' in nf and len(nf) <= len(_norm('Name:') + 'softwarecomponentversion')
                if is_name_row and curr:
                    looks_like_wrong_name = (
                        curr.startswith('MM_') or
                        curr.startswith('GLO') or
                        (len(curr) > 5 and not curr.startswith('[') and
                         not curr.startswith('http') and
                         not curr.startswith('This') and
                         any(sep in curr for sep in ('_', '-')) and
                         curr != mmap_name)
                    )
                    if looks_like_wrong_name:
                        _set_cell(val_cell, mmap_name)

            # ── Artifact: <name> (CPI) rows ───────────────────────────────────
            elif mmap_name and 'artifactnamespace' in nf or ('artifact' in nf and 'cpi' in nf and 'namespace' not in nf):
                if curr and curr != f'Artifact: {mmap_name} (CPI)' and 'Artifact:' in curr:
                    _set_cell(val_cell, f'Artifact: {mmap_name} (CPI)')


def update_td_with_iflow(td_bytes: bytes, iflow_zip_bytes: bytes, author: str = '') -> bytes:
    """
    Update an existing TD document with iFlow ZIP data.
    - Fills main body tables from Appendix data
    - Overrides iFlow artifact fields from ZIP MANIFEST
    - Adds iFlow Design Steps section
    - Adds SAP-themed flow diagram
    ZERO AI — 100% deterministic.
    """
    doc   = Document(io.BytesIO(td_bytes))
    facts = _extract_iflow_facts(iflow_zip_bytes)
    # Overlay adapter and step data from parse_iflow_zip which correctly
    # maps participant IDs → names (source_name/target_name) and infers direction.
    # _extract_iflow_facts is kept for mmap/xsd/script fields.
    try:
        _parsed = _parse_iflow_zip_full(iflow_zip_bytes)
        if _parsed.get('adapters'):
            facts['adapters'] = _parsed['adapters']
        if _parsed.get('steps'):
            facts['steps'] = _parsed['steps']
    except Exception:
        pass  # fall back to _extract_iflow_facts data if parse fails
    appendix_data = _extract_appendix_data(doc)

    # ── Step 0: Extract what the TD currently has (to know what to replace) ──
    # Collect all "wrong" artifact names already in the document that need replacing.
    # We collect every cell value in the iFlow artifact and mapping tables
    # and replace any that differ from the correct ZIP names.
    existing_iflow_names = set()
    existing_mmap_names  = set()

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if not cells: continue
            for cell in cells:
                t = cell.text.strip()
                # Looks like an iFlow artifact name (CamelCase, no spaces, starts with Send/Receive/Map etc.)
                if (len(t) > 8 and ' ' not in t and
                    any(t.startswith(p) for p in ('Send','Receive','Get','Post','Map','I_','GLO','MM_',
                                                    'FDMAP','TDMAP','Inbound','Outbound'))):
                    # It's a candidate wrong name if it's NOT the correct names
                    if t != facts['iflow_name'] and t != facts['mmap_name']:
                        if any(kw in cell.text for kw in ['Batch','Stock','MBGMCR','ITALTRANS','Italtrans',
                                                            'BatchStatus','StockStatus']):
                            if t.startswith('MM_'):
                                existing_mmap_names.add(t)
                            else:
                                existing_iflow_names.add(t)

    # ── Step 0b: Build replacement list and do a full document text replacement
    replacements = []
    iflow_name = facts.get('iflow_name', '')
    mmap_name  = facts.get('mmap_name', '')

    # Replace wrong iFlow names with correct one
    if iflow_name:
        for wrong in existing_iflow_names:
            if wrong and wrong != iflow_name:
                replacements.append((wrong, iflow_name))
        # Also replace "Artifact: {wrong} (CPI)" → "Artifact: {iflow_name} (CPI)"
        for wrong in existing_iflow_names:
            replacements.append((f'Artifact: {wrong} (CPI)', f'Artifact: {iflow_name} (CPI)'))

    # Replace wrong mmap names with correct one
    if mmap_name:
        for wrong in existing_mmap_names:
            if wrong and wrong != mmap_name:
                replacements.append((wrong, mmap_name))

    if replacements:
        _replace_all_occurrences(doc, replacements)

    # ── Step 0c: Targeted fix for Mapping table ───────────────────────────────
    # After global replacement, the Mapping table rows (16) may now have the
    # iFlow name where they should have the mmap name.
    # Fix: in any table with a "Mapping" header, replace iFlow name → mmap name.
    if iflow_name and mmap_name:
        for table in doc.tables:
            # Check if this table has a "Mapping" header
            first_row_text = ' '.join(c.text.strip() for c in table.rows[0].cells) if table.rows else ''
            if 'Mapping' not in first_row_text:
                continue
            # In this mapping table, replace iFlow name with mmap name
            fix_pairs = [
                (iflow_name, mmap_name),
                (f'Artifact: {iflow_name} (CPI)', f'Artifact: {mmap_name} (CPI)'),
                (f'Artifact: {iflow_name}', f'Artifact: {mmap_name}'),
            ]
            for row in table.rows:
                for cell in row.cells:
                    for old_t, new_t in fix_pairs:
                        if old_t in cell.text:
                            _replace_text_in_cell(cell, old_t, new_t)

    # ── Build override dict from actual iFlow ZIP facts ───────────────────────
    overrides = {}
    if facts['iflow_name']:
        overrides[_norm('IFlow (PO) or Artifact')] = facts['iflow_name']
        overrides[_norm('Artifact (CPI)')] = facts['iflow_name']
        # Working Name of Interface should be the iFlow name, not the mmap name
        overrides[_norm('Working Name of Interface')] = facts['iflow_name']
    if facts['description']:
        overrides[_norm('Description')] = facts['description']
    # NOTE: 'Name:' override for mmap deliberately removed — it matched too broadly
    # (e.g. 'name' is a substring of 'nameofapplication', 'workingnameofinterface').
    # Mmap name in Mapping tables is handled by the dedicated targeted fix below.
    if facts['iflow_version']:
        overrides[_norm('Version')] = facts['iflow_version']

    # ── Full-document name replacement ────────────────────────────────────────
    # BEFORE filling tables, do a direct find-and-replace of wrong names
    # across ALL cells in ALL tables. This catches names that were copied
    # from the FD into the TD and are now incorrect.
    _replace_artifact_names_everywhere(doc, facts, appendix_data)

    # ── Fill each main body table from appendix data + overrides ─────────────
    # Appendix tables are the last ~14 tables; main body is everything before
    n_tables = len(doc.tables)
    main_body_limit = max(0, n_tables - 14)

    for i, table in enumerate(doc.tables[:main_body_limit]):
        _fill_table_from_appendix(table, appendix_data, overrides)

    # ── Fill author name if provided ──────────────────────────────────────────
    if author and author.strip():
        for table in doc.tables[:main_body_limit]:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2 and 'Author' in cells[0].text and len(cells[0].text.strip()) < 10:
                    val_cell = cells[1]
                    if val_cell != cells[0]:
                        _set_cell(val_cell, author.strip())
                    break

    # NOTE: Integration Flow step summary intentionally NOT injected —
    # the existing TD content for "Integration flow:" is preserved as-is.

    # ── Hard replace: iFlow name, package, mapping name in ALL tables ────────
    # These must be replaced regardless of current cell content because
    # the TD may have wrong names from FD (e.g. SendBatchAndStockStatusToItaltrans
    # instead of the real iFlow Bundle-Name from the ZIP).
    iflow_name_from_zip = facts['iflow_name']
    mmap_name_from_zip  = facts['mmap_name']
    # Pass appendix text so country detection finds "Italy" from "Italy Italtrans 3PL"
    _appendix_text = ' '.join(str(v) for v in appendix_data.values())
    pkg_from_appendix   = derive_package_name(facts, extra_text=_appendix_text)

    for table in doc.tables[:main_body_limit]:
        for row in table.rows:
            cells = row.cells
            if not cells: continue
            first = cells[0].text.strip()
            if not first: continue
            nf = _norm(first)

            # Get the value cell (first non-merged cell after label)
            val_cell = None
            for c in cells[1:]:
                if c.text.strip() != cells[0].text.strip():
                    val_cell = c
                    break
            if not val_cell:
                continue

            # Package / Folder name → ALWAYS replace with appendix value
            if ('folder' in nf or 'package' in nf) and pkg_from_appendix:
                _set_cell(val_cell, pkg_from_appendix)

            # iFlow / Artifact name → ALWAYS replace with ZIP Bundle-Name
            elif ('iflow' in nf or ('artifact' in nf and 'cpi' in nf)) and iflow_name_from_zip:
                _set_cell(val_cell, iflow_name_from_zip)

            # iFlow Description → ALWAYS replace with metainfo.prop description
            elif nf == _norm('Description:') and facts['description']:
                _set_cell(val_cell, facts['description'])

            # Mapping Name rows → ALWAYS replace with actual mmap name from ZIP
            elif mmap_name_from_zip and val_cell.text.strip():
                # If the cell looks like an old mapping name (starts with MM_ or SendBatch...)
                curr = val_cell.text.strip()
                is_old_mapping_name = (
                    (curr.startswith('MM_') or curr.startswith('Send') or
                     curr.startswith('FDMAP') or curr.startswith('TDMAP'))
                    and 'Artifact:' not in curr
                    and any(k in nf for k in ('name', 'artifact', 'mapping'))
                )
                if is_old_mapping_name:
                    # Only replace if the label suggests it's a mapping name field
                    if 'name' in nf and len(nf) < 15:
                        _set_cell(val_cell, mmap_name_from_zip)

    # ── Step FINAL: Re-apply Mapping table fix after all other fills ─────────
    # The appendix fill / hard-replace may have re-set iFlow name into Mapping
    # table cells. Run the targeted replacement one final time as the last op
    # so the mmap name is authoritative in all Mapping tables.
    if iflow_name_from_zip and mmap_name_from_zip:
        fix_pairs_final = [
            (iflow_name_from_zip, mmap_name_from_zip),
            (f'Artifact: {iflow_name_from_zip} (CPI)', f'Artifact: {mmap_name_from_zip} (CPI)'),
            (f'Artifact: {iflow_name_from_zip}', f'Artifact: {mmap_name_from_zip}'),
        ]
        for table in doc.tables:
            first_row_text = ' '.join(c.text.strip() for c in table.rows[0].cells) if table.rows else ''
            if 'Mapping' not in first_row_text:
                continue
            for row in table.rows:
                for cell in row.cells:
                    for old_t, new_t in fix_pairs_final:
                        if old_t in cell.text:
                            _replace_text_in_cell(cell, old_t, new_t)

    # ── Tick all relevant checkboxes from iFlow/appendix data ─────────────────
    _all_adapters = facts.get('adapters', [])

    # ① Middleware Solutions — CPI always; Event Mesh if SAP AEM sender
    _has_aem = any(
        any(k in (a.get('source_name', '') + a.get('name', '')).lower()
            for k in ('aem', 'event', 'mesh'))
        for a in _all_adapters
    )
    _cb_to_check = {'CPI'}
    if _has_aem:
        _cb_to_check.add('Event Mesh')

    # ② Business Criticality — read from appendix data
    _biz_crit_val = ''
    for _ak, _av in appendix_data.items():
        if 'businesscriticality' in _ak or ('business' in _ak and 'critica' in _ak):
            _biz_crit_val = str(_av).strip().lower()
            break
    if 'highly' in _biz_crit_val:
        _cb_to_check.add('Highly critical')
    elif 'non' in _biz_crit_val or 'not' in _biz_crit_val:
        _cb_to_check.add('Non-critical')
    elif 'critical' in _biz_crit_val:
        _cb_to_check.add('Critical')

    # ③ Data flow direction — Unidirectional / Multiple from adapter count
    _real_receivers = [a for a in _all_adapters
                       if a.get('direction') == 'Receiver'
                       and a.get('component') not in ('ProcessDirect',)]
    if len(_real_receivers) > 1:
        _cb_to_check.add('Multiple')
    else:
        _cb_to_check.add('Unidirectional')

    try:
        _check_middleware_checkboxes(doc, _cb_to_check)
    except Exception:
        pass  # best-effort

    # ── Fix "Within Sender" monitoring section ────────────────────────────────
    _real_sender = appendix_data.get(_norm('Source system'), '') or 'S/4HANA'
    _is_s4_sender = any(k in _real_sender.upper() for k in ('S4', 'S/4', 'HANA', 'ECC', 'ERP'))
    if _is_s4_sender:
        for table in doc.tables[:main_body_limit]:
            header_text = table.rows[0].cells[0].text.strip() if table.rows else ''
            if 'Within Sender' in header_text:
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        label = cells[0].text.strip()
                        val   = cells[1].text.strip()
                        if 'Monitoring' in label and val in ('Third Party System', 'Third party System', 'Third Party'):
                            _set_cell(cells[1], f'SAP {_real_sender} system monitoring (Transaction SXMB_MONI / Cloud ALM)')
                break

    # ── Fix Source/Target "Name of application" + Type of System + Protocol ──
    _src_app = appendix_data.get(_norm('Source system'), '') or _real_sender

    # Use adapter target_name directly — avoids wrong participant names like
    # "TriggerException" that _extract_vendor_name may return
    _real_recv_adapters = [a for a in _all_adapters
                           if a.get('direction') == 'Receiver'
                           and a.get('component') not in ('ProcessDirect',)]
    _tgt_app = (
        (_real_recv_adapters[0].get('target_name') or _real_recv_adapters[0].get('name', ''))
        if _real_recv_adapters
        else (appendix_data.get(_norm('Target system'), '') or 'Target')
    )

    # Adapter protocols for checkbox labelling
    _sender_proto  = (_all_adapters[next((i for i, a in enumerate(_all_adapters)
                                          if a.get('direction') == 'Sender'), -1)]
                      .get('component', 'HTTPS') if any(a.get('direction') == 'Sender'
                                                         for a in _all_adapters) else 'HTTPS')
    _target_proto  = (_real_recv_adapters[0].get('component', '') if _real_recv_adapters else '')

    # Determine system types
    _src_is_sap = any(k in _src_app.upper() for k in ('S4', 'S/4', 'HANA', 'ECC', 'ERP', 'SAP'))
    _tgt_is_sap = any(k in _tgt_app.upper() for k in ('S4', 'S/4', 'HANA', 'ECC', 'ERP', 'SAP'))

    for _tbl in doc.tables[:main_body_limit]:
        _hdr = _tbl.rows[0].cells[0].text.strip() if _tbl.rows else ''

        if _hdr in ('Source', 'Source System'):
            # Name of application
            for _row in _tbl.rows[1:]:
                if len(_row.cells) >= 2 and 'name of application' in _row.cells[0].text.lower():
                    _set_cell(_row.cells[1], _src_app)
                    break
            # Checkboxes: Type of System + Communication Protocol
            try:
                _src_cb = set()
                if _src_is_sap:
                    _src_cb.add('SAP System')
                else:
                    _src_cb.add('Other (including non-SAP')
                if _sender_proto:
                    _src_cb.add(_sender_proto)   # e.g. 'HTTPS', 'AMQP'
                if _has_aem:
                    _src_cb.add('AMQP')          # S/4→AEM uses AMQP
                if _src_cb:
                    _check_middleware_checkboxes(_tbl._tbl, _src_cb)
            except Exception:
                pass

        elif _hdr in ('Target', 'Target System'):
            # Name of application
            for _row in _tbl.rows[1:]:
                if len(_row.cells) >= 2 and 'name of application' in _row.cells[0].text.lower():
                    _set_cell(_row.cells[1], _tgt_app)
                    break
            # Checkboxes: Type of System + Communication Protocol
            try:
                _tgt_cb = set()
                if _tgt_is_sap:
                    _tgt_cb.add('SAP System')
                else:
                    _tgt_cb.add('Third-Party System')
                    _tgt_cb.add('Other (including non-SAP')  # some TDs use this label
                if _target_proto:
                    _tgt_cb.add(_target_proto)   # e.g. 'AS2', 'SFTP'
                if _tgt_cb:
                    _check_middleware_checkboxes(_tbl._tbl, _tgt_cb)
            except Exception:
                pass

    # ── Embed mapping Excel OLE into the existing Mapping Objects table ───────
    # Generates the xlsx, then inserts an OLE paragraph directly after the
    # Mapping table (section 3.1.2) in the main body using XML manipulation.
    if facts.get('mmap_name'):
        try:
            from services.mapping_excel import generate_mapping_excel
            from lxml import etree as _ET
            import uuid as _uuid
            _xlsx_result = generate_mapping_excel(iflow_zip_bytes)
            if _xlsx_result:
                _xlsx_bytes, _xlsx_embed_name = _xlsx_result
                _rId_obj  = f"rIdMmap{_uuid.uuid4().hex[:8]}"
                _shape_id = f"_x0000_i{_uuid.uuid4().int % 90000 + 10000}"
                _obj_id   = f"_{_uuid.uuid4().int % 2000000000}"
                _NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                _NS_V = 'urn:schemas-microsoft-com:vml'
                _NS_O = 'urn:schemas-microsoft-com:office:office'
                _NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

                # Build OLE paragraph XML.
                # IMPORTANT: No <v:imagedata> — a PNG imagedata makes Word treat the
                # whole thing as a picture (Picture Format ribbon) instead of an OLE
                # object.  Use a solid-filled rectangle (type t1) with o:ole="" so
                # Word recognises it as an OLE container and allows double-click.
                _short_name = _xlsx_embed_name.replace('_MappingSpec.xlsx', '')
                _ole_xml = (
                    f'<w:p xmlns:w="{_NS_W}" xmlns:v="{_NS_V}" '
                    f'xmlns:o="{_NS_O}" xmlns:r="{_NS_R}">'
                    f'<w:pPr><w:jc w:val="left"/></w:pPr>'
                    f'<w:r><w:object w:dxaOrig="4320" w:dyaOrig="1008">'
                    # type #_x0000_t1 = plain rectangle — no imagedata needed,
                    # solid green fill, white border, OLE container marker
                    f'<v:shape id="{_shape_id}" type="#_x0000_t1" '
                    f'fillcolor="#217346" strokecolor="#FFFFFF" strokeweight="1pt" '
                    f'style="width:3in;height:0.7in" o:ole="">'
                    f'<v:fill type="solid" color="#217346"/>'
                    f'<v:stroke color="#FFFFFF" weight="1pt"/>'
                    f'<v:textbox style="mso-direction-alt:auto;mso-fit-shape-to-text:false" '
                    f'inset="4pt,4pt,4pt,4pt">'
                    f'<w:txbxContent>'
                    f'<w:p><w:pPr><w:jc w:val="left"/></w:pPr>'
                    f'<w:r><w:rPr><w:b/><w:color w:val="FFFFFF"/>'
                    f'<w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>'
                    f'<w:t xml:space="preserve">■ {_short_name}.xlsx</w:t></w:r></w:p>'
                    f'<w:p><w:r><w:rPr><w:color w:val="C8FFD4"/>'
                    f'<w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr>'
                    f'<w:t>Double-click to open mapping specification</w:t>'
                    f'</w:r></w:p>'
                    f'</w:txbxContent></v:textbox>'
                    f'</v:shape>'
                    f'<o:OLEObject Type="Embed" ProgID="Excel.Sheet.12" '
                    f'ShapeID="{_shape_id}" DrawAspect="Icon" '
                    f'ObjectID="{_obj_id}" r:id="{_rId_obj}"/>'
                    f'</w:object></w:r>'
                    f'</w:p>'
                )
                _ole_elem = _ET.fromstring(_ole_xml)

                # Find the FIRST Mapping table in main body and insert after it
                _inserted = False
                for _mt in doc.tables[:main_body_limit]:
                    _mhdr = _mt.rows[0].cells[0].text.strip() if _mt.rows else ''
                    if 'Mapping' in _mhdr and len(_mhdr) < 20:
                        _mt._tbl.addnext(_ole_elem)
                        _inserted = True
                        break

                _ole_ids = {
                    'rId_obj':    _rId_obj,
                    'embed_name': _xlsx_embed_name,
                }
        except Exception as _exc:
            pass  # silent — OLE is best-effort

    # ── Add new page with iFlow Design Steps ─────────────────────────────────
    doc.add_page_break()

    h = doc.add_heading('iFlow Design – Technical Steps', level=1)
    for r in h.runs: r.font.color.rgb = SAP_BLUE

    # Use iflow_name (Bundle-Name) which is now set by parse_iflow_zip
    iflow_display_name = facts.get('iflow_name') or facts.get('name', 'Integration Process')
    pkg_name = derive_package_name(facts, extra_text=_appendix_text)

    _add_para(doc,
        f'iFlow: {iflow_display_name} | Version: {facts.get("iflow_version","1.0")} | '
        f'Package: {pkg_name}',
        size=9, color=MID_GREY)
    _add_para(doc, facts['description'], size=10)
    doc.add_paragraph()

    # ── Flow diagram ─────────────────────────────────────────────────────────
    h2 = doc.add_heading('Integration Flow Diagram', level=2)
    for r in h2.runs: r.font.color.rgb = SAP_BLUE

    senders   = [a for a in facts['adapters'] if a['direction'] == 'Sender']
    receivers = [a for a in facts['adapters'] if a['direction'] == 'Receiver']
    main_steps= [s for s in facts['steps']
                 if s['element_type'] not in ('endEvent',)
                 and 'Error' not in s.get('name','')
                 and 'Error' not in s.get('element_type','')]

    step_desc = ', '.join(f"{s['name']} [{STEP_LABELS.get(s['activity_type'], s['element_type'])}]"
                          for s in main_steps if s.get('name'))
    act_types = [s['activity_type'] for s in main_steps]
    mapping_type = ('Groovy' if any(a in ('Script','GroovyScript') for a in act_types)
                    else 'Message Mapping' if 'Mapping' in act_types else '')

    # Adapters from parse_iflow_zip use source_name/target_name (not source_ref/target_ref)
    sender_name = (senders[0].get('source_name') or senders[0].get('name','')
                   or 'Source System') if senders else 'Source System'
    sender_comp = senders[0]['component'] if senders else 'HTTPS'

    # Real receivers = exclude ProcessDirect (internal, not external target)
    real_receivers = [r for r in receivers if r.get('component') not in ('ProcessDirect',)]
    # Use only the system name (not the adapter protocol) in the diagram box label
    tgt_display = ', '.join(
        (r.get('target_name') or r.get('name','') or 'Target').strip()
        for r in real_receivers
    ) if real_receivers else 'Target System'
    tgt_protocol = real_receivers[0]['component'] if real_receivers else 'HTTP'

    # ── Build full process chain for diagram ──────────────────────────────────
    # If the sender is SAP AEM (Event Mesh), the real end-to-end flow is:
    #   S/4HANA → AEM → CPI → Target
    # Detect this and prepend the upstream source system so the diagram shows
    # the complete chain, not just AEM → CPI → Target.
    _is_aem_sender = any(k in sender_name.lower() for k in ('aem', 'event', 'mesh'))
    if _is_aem_sender:
        # Try to get upstream source system from TD appendix data
        _upstream = (appendix_data.get(_norm('Source system'), '')
                     or appendix_data.get(_norm('Source system:'), '')
                     or 'S/4HANA')
        # Build explicit chain: S/4HANA → AEM → CPI → Target
        process_flow_chain = f'{_upstream} → {sender_name} → CPI → {tgt_display}'
    else:
        # Single sender, let _parse_chain fall back to source_app_name / target_app_name
        process_flow_chain = ''

    correct_png = None
    try:
        correct_png = generate_flowchart({
            'interface_name':  iflow_display_name,
            'source_app_name': sender_name,
            'source_protocol': sender_comp,
            'target_app_name': tgt_display,
            'target_protocol': tgt_protocol,
            'process_flow':    process_flow_chain,   # full S4→AEM→CPI→Target chain
            'show_steps':      False,                # clean CPI box, no internal steps
        })
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(correct_png), width=Inches(6.2))
    except Exception as e:
        _add_para(doc, f'[Diagram error: {e}]', size=9, color=RED_C)

    doc.add_paragraph()

    # ── Narrative step-by-step guide ──────────────────────────────────────────
    _build_narrative_steps(doc, facts, senders, receivers)

    # ── Groovy Scripts ────────────────────────────────────────────────────────
    if facts['scripts']:
        h7 = doc.add_heading('Groovy Scripts', level=2)
        for r in h7.runs: r.font.color.rgb = SAP_BLUE
        for fname, code in facts['scripts'].items():
            t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
            c = t.rows[0].cells[0]; c.text = fname
            _shd(c, '1A1A2E')
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            c.paragraphs[0].runs[0].font.size = Pt(10)
            _add_code(doc, code[:3000])
            doc.add_paragraph()

    # ── XSD / Schema references ───────────────────────────────────────────────
    if facts['xsd_files']:
        h8 = doc.add_heading('XSD Schema Files', level=2)
        for r in h8.runs: r.font.color.rgb = SAP_BLUE
        for xf in facts['xsd_files']:
            _add_para(doc, f'• {xf}', size=10)
        doc.add_paragraph()

    # ── Save ─────────────────────────────────────────────────────────────────
    out = io.BytesIO()
    doc.save(out)
    raw = out.getvalue()

    # ── ZIP surgery: replace diagrams + embed mapping Excel OLE ─────────────
    if correct_png or _ole_ids:
        try:
            import zipfile as _zf
            buf2 = io.BytesIO()
            with _zf.ZipFile(io.BytesIO(raw), 'r') as zin, \
                 _zf.ZipFile(buf2, 'w', compression=_zf.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    fn   = item.filename

                    # ① Replace stale flowchart PNGs with correct diagram
                    if (correct_png and fn.startswith('word/media/') and
                            fn.lower().endswith('.png') and len(data) > 30000):
                        data = correct_png

                    # ② Inject OLE relationship for embedded Excel
                    elif _ole_ids and fn == 'word/_rels/document.xml.rels':
                        data = data.decode('utf-8')
                        new_rel = (
                            f'<Relationship Id="{_ole_ids["rId_obj"]}" '
                            f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" '
                            f'Target="embeddings/{_ole_ids["embed_name"]}"/>'
                        )
                        data = data.replace('</Relationships>', new_rel + '</Relationships>')
                        data = data.encode('utf-8')

                    # ③ Register xlsx content-type so Word recognises it
                    elif _ole_ids and fn == '[Content_Types].xml':
                        data = data.decode('utf-8')
                        ct = ('application/vnd.openxmlformats-officedocument'
                              '.spreadsheetml.sheet')
                        if 'xlsx' not in data:
                            data = data.replace('</Types>',
                                f'<Default Extension="xlsx" ContentType="{ct}"/>'
                                '</Types>')
                        data = data.encode('utf-8')

                    zout.writestr(item, data)

                # ④ Write embedded xlsx file into the docx package
                if _ole_ids and _xlsx_bytes:
                    zout.writestr(
                        f'word/embeddings/{_ole_ids["embed_name"]}',
                        _xlsx_bytes,
                    )

            return buf2.getvalue()
        except Exception:
            pass  # fall back to unmodified saved bytes

    return raw
