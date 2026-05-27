"""
doc_builder.py
Builds .docx FD and TD documents.
TD format matches the reference: TDMAP_I_001MaterialmasteroutfromSAPsystemtoExternalsystem1_TD.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from io import BytesIO
from datetime import datetime
from services.flowchart_builder import generate_flowchart


# ── colour palette ────────────────────────────────────────────────────────────
SAP_BLUE   = "0070F2"
LIGHT_BLUE = "DDEEFF"
GRAY_ROW   = "F5F5F5"
WHITE      = "FFFFFF"

# TD-specific colours (matched to reference document)
HDR_DARK   = "17375E"   # dark navy — spanning section headers (Source, Target, Within Sender…)
HDR_MED    = "2E74B5"   # medium blue — column headers in data tables
LBL_BG     = "DEEAF1"   # light blue — label column in key-value tables
ROW_ALT    = "F2F2F2"   # very light grey — alternate body rows


# ── low-level XML helpers ─────────────────────────────────────────────────────

def _shd(hex_color: str):
    return parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )


def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove any existing shading first
    for existing in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(existing)
    tc_pr.append(_shd(hex_color))


def _cell_text(cell, text: str, bold: bool = False,
               fg: str = "000000", size: int = 9, italic: bool = False):
    """Set cell text, clearing prior content."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(fg)


def _hdr_cell(cell, text: str, bg: str = HDR_DARK):
    """Dark/coloured spanning header cell — white bold text."""
    _set_cell_bg(cell, bg)
    _cell_text(cell, text, bold=True, fg="FFFFFF", size=9)


def _lbl_cell(cell, text: str):
    """Label cell — light-blue background, dark text."""
    _set_cell_bg(cell, LBL_BG)
    _cell_text(cell, text, bold=True, fg="000000", size=9)


def _val_cell(cell, text: str, alt: bool = False):
    """Value cell — white (or alt grey) background."""
    _set_cell_bg(cell, ROW_ALT if alt else WHITE)
    _cell_text(cell, text or "", fg="000000", size=9)


def _col_hdr_cell(cell, text: str):
    """Column header cell for data tables — medium blue, white bold."""
    _set_cell_bg(cell, HDR_MED)
    _cell_text(cell, text, bold=True, fg="FFFFFF", size=9)


def _merge_row(row, num_cols: int):
    """Merge all cells in a row into one."""
    if num_cols < 2:
        return row.cells[0]
    merged = row.cells[0]
    for i in range(1, num_cols):
        merged = merged.merge(row.cells[i])
    return merged


def _set_col_width(table, widths: list):
    """Set column widths (in Inches). widths is a list per column."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])


# ── mid-level table builders ──────────────────────────────────────────────────

def _spanning_header_table(doc: Document, header: str, rows: list,
                            col_widths=(2.2, 4.1)):
    """
    Two-column key-value table with a dark spanning header row.
    rows: list of (label_str, value_str)
    This is the primary table style in the reference document:
      ┌────────────────────────────────────────┐
      │  Header (dark navy, white bold)        │
      ├─────────────────────┬──────────────────┤
      │ Label (light blue)  │ Value (white)    │
      └─────────────────────┴──────────────────┘
    """
    num_rows = 1 + len(rows)
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = "Table Grid"

    # Spanning header
    hdr_cell = _merge_row(table.rows[0], 2)
    _hdr_cell(hdr_cell, header)

    # Label-value rows
    for i, (label, value) in enumerate(rows):
        r = table.rows[i + 1]
        _lbl_cell(r.cells[0], label)
        _val_cell(r.cells[1], value, alt=(i % 2 == 1))

    _set_col_width(table, list(col_widths))
    doc.add_paragraph()
    return table


def _plain_kv_table(doc: Document, rows: list, col_widths=(2.2, 4.1)):
    """
    Two-column key-value table WITHOUT a spanning header.
    rows: list of (label_str, value_str)
    """
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        _lbl_cell(table.rows[i].cells[0], label)
        _val_cell(table.rows[i].cells[1], value, alt=(i % 2 == 1))
    _set_col_width(table, list(col_widths))
    doc.add_paragraph()
    return table


def _data_table(doc: Document, headers: list, rows: list):
    """
    N-column data table with a medium-blue header row.
    rows: list of tuples/lists of strings.
    """
    n_cols = len(headers)
    n_rows = 1 + len(rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        _col_hdr_cell(table.rows[0].cells[j], h)

    for i, row_data in enumerate(rows):
        bg = ROW_ALT if i % 2 == 1 else WHITE
        for j in range(n_cols):
            val = row_data[j] if j < len(row_data) else ""
            _set_cell_bg(table.rows[i + 1].cells[j], bg)
            _cell_text(table.rows[i + 1].cells[j], val, size=9)

    doc.add_paragraph()
    return table


def _bullet(doc: Document, text: str):
    doc.add_paragraph(str(text), style="List Bullet")


def _bold_para(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.bold = True
    run.font.size = Pt(10)


def _h(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def _ensure_list(val) -> list:
    if isinstance(val, list):
        return [str(v) for v in val if v]
    if val and str(val).strip() not in ("", "N/A", "None"):
        return [str(val)]
    return []


def _join_list(val) -> str:
    if isinstance(val, list):
        return ", ".join(str(v) for v in val if v)
    return str(val) if val else ""


def _safe(val, default: str = "") -> str:
    """Return the value as a string, or default (empty string) if absent/None."""
    if val is None:
        return default
    s = str(val).strip()
    return s if s else default


# ── Table of Contents ─────────────────────────────────────────────────────────

def _add_toc(doc: Document):
    """
    Insert a Word TOC field (\\o 1-3, hyperlinked).
    When the .docx is opened in Word, right-click the grey area
    and choose 'Update Field' (or press Ctrl+A, F9) to render the TOC.
    """
    doc.add_heading("Contents", level=1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # --- begin field ---
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    # --- field instruction ---
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    # --- separator + placeholder text ---
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_sep)

    ph_run = OxmlElement('w:r')
    ph_rpr = OxmlElement('w:rPr')
    ph_nopr = OxmlElement('w:noProof')
    ph_rpr.append(ph_nopr)
    ph_run.append(ph_rpr)
    ph_t = OxmlElement('w:t')
    ph_t.text = (
        '[Open in Word and press Ctrl+A then F9 to update the Table of Contents]'
    )
    ph_run.append(ph_t)
    run._r.append(ph_run)

    # --- end field ---
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)

    doc.add_page_break()


# ── FD Builder (unchanged) ─────────────────────────────────────────────────────

def _bold_cell(cell, text: str, bg: str = SAP_BLUE, color: str = "FFFFFF", size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    _set_cell_bg(cell, bg)


def _normal_cell(cell, text: str, bg: str = WHITE, size: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text else "N/A")
    run.font.size = Pt(size)
    _set_cell_bg(cell, bg)


def _add_heading(doc: Document, text: str, level: int):
    doc.add_heading(text, level=level)


def _add_kv_table(doc: Document, rows: list):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        bg = GRAY_ROW if i % 2 == 0 else WHITE
        _bold_cell(table.rows[i].cells[0], key, bg=LIGHT_BLUE, color="000000")
        _normal_cell(table.rows[i].cells[1], val, bg=bg)
        table.rows[i].cells[0].width = Inches(2.2)
        table.rows[i].cells[1].width = Inches(4.0)
    doc.add_paragraph()


def _add_section_table(doc: Document, headers: list, data: list):
    cols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=cols)
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        _bold_cell(table.rows[0].cells[j], h)
    for i, row in enumerate(data):
        bg = GRAY_ROW if i % 2 == 0 else WHITE
        for j, val in enumerate(row):
            _normal_cell(table.rows[i + 1].cells[j], val, bg=bg)
    doc.add_paragraph()


def build_fd(data: dict) -> bytes:
    data.setdefault("source_app_name", data.get("from_system", ""))
    data.setdefault("target_app_name", data.get("to_system", ""))
    data.setdefault("source_protocol", data.get("processing_type", ""))
    data.setdefault("target_protocol", data.get("processing_type", ""))
    data.setdefault("integration_logic", data.get("process_flow", ""))

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Functional Design – Interface")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(SAP_BLUE)

    doc.add_paragraph(f"[{data.get('interface_id', '')}]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(data.get("interface_name", "")).runs[0].bold = True
    doc.add_paragraph()

    _add_heading(doc, "Document Control Information", 1)
    _add_heading(doc, "Document Revision History", 2)
    _add_section_table(doc,
        ["Version", "Date", "Prepared By", "Description"],
        [["V1.0", datetime.today().strftime("%Y-%m-%d"), data.get("author", ""), "Initial Draft"]])

    _add_heading(doc, "Overview", 1)
    doc.add_paragraph(data.get("overview", ""))
    _add_heading(doc, "To-Be Process", 2)
    doc.add_paragraph(data.get("to_be_process", ""))
    _add_heading(doc, "Object Overview", 2)
    doc.add_paragraph(data.get("object_overview", ""))
    _add_heading(doc, "Business Details", 2)
    doc.add_paragraph(data.get("business_details", ""))

    _add_heading(doc, "Overview and Scope", 1)
    _add_heading(doc, "Functional Description", 2)
    doc.add_paragraph(data.get("functional_description", ""))
    _add_heading(doc, "Interface From – To", 2)
    _add_kv_table(doc, [
        ("From System", data.get("from_system", "")),
        ("Transformation System", data.get("transformation_system", "SAP CPI")),
        ("To System", data.get("to_system", "")),
        ("Processing Type", data.get("processing_type", "")),
    ])
    _add_heading(doc, "Assumptions", 2)
    doc.add_paragraph(data.get("assumptions", "N/A"))
    _add_heading(doc, "Dependencies / Constraints", 2)
    doc.add_paragraph(data.get("dependencies", "N/A"))
    _add_heading(doc, "Security, Integrity and Controls", 2)
    doc.add_paragraph(data.get("security", ""))
    _add_heading(doc, "Fiori Impact Assessment", 2)
    doc.add_paragraph(data.get("fiori_impact", "N/A"))

    _add_heading(doc, "Detailed Functional Requirements", 1)
    _add_heading(doc, "Functional Specification", 2)
    doc.add_paragraph(data.get("functional_spec", ""))
    _add_heading(doc, "Interface Process Flow Diagram", 2)
    doc.add_paragraph(data.get("process_flow", ""))
    try:
        png_bytes = generate_flowchart(data)
        doc.add_picture(BytesIO(png_bytes), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Flowchart could not be generated: {e}]")
    doc.add_paragraph()

    _add_heading(doc, "Proposed Transfer Method / Program", 2)
    doc.add_paragraph(data.get("transfer_method", ""))
    _add_heading(doc, "Target Data Layout", 2)
    doc.add_paragraph(data.get("target_data_layout", ""))
    _add_heading(doc, "Source Data Layout", 2)
    doc.add_paragraph(data.get("source_data_layout", ""))
    _add_heading(doc, "Mapping SAP Fields to Source/Target", 2)
    if data.get("field_mappings"):
        _add_section_table(doc,
            ["Source Field", "Source Path", "Target Field", "Transformation"],
            data["field_mappings"])
    else:
        doc.add_paragraph("Refer to attached mapping file.")

    _add_heading(doc, "Post Interface Activities", 2)
    doc.add_paragraph(data.get("post_activities", "N/A"))
    _add_heading(doc, "Reporting / Monitoring", 2)
    doc.add_paragraph(data.get("monitoring", ""))

    _add_heading(doc, "Additional Requirements", 1)
    _add_heading(doc, "Authorization Checks", 2)
    doc.add_paragraph(data.get("authorization", "N/A"))
    _add_heading(doc, "Messages / Error Handling", 2)
    doc.add_paragraph(data.get("error_handling", ""))
    _add_heading(doc, "Performance Considerations", 2)
    doc.add_paragraph(data.get("performance", "N/A"))

    _add_heading(doc, "Testing Requirements", 1)
    _add_heading(doc, "Key Business Test Conditions", 2)
    for t in data.get("business_test_conditions", []):
        doc.add_paragraph(t, style="List Bullet")
    _add_heading(doc, "Technical Test Conditions", 2)
    for t in data.get("technical_test_conditions", []):
        doc.add_paragraph(t, style="List Bullet")

    _add_heading(doc, "Outstanding Issues", 1)
    issues = data.get("issues", [])
    if issues:
        _add_section_table(doc, ["Issue", "Status", "Owner"], issues)
    else:
        doc.add_paragraph("N/A")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── TD Builder — exact match of reference document format ─────────────────────

def _td_footer(doc: Document):
    """Add 'Classification: Internal' to all section footers."""
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.clear()
        run = p.add_run("Classification: Internal")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _cover_page(doc: Document, data: dict):
    """
    Cover page matching the reference:
    - Large title 'Integration Technical Specification'
    - Version History table
    - Document Information table
    """
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run("Integration Technical Specification")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor.from_string(HDR_DARK)

    doc.add_paragraph()

    # Version History
    _bold_para(doc, "Version History")
    _data_table(doc,
        ["Ver.", "Date", "Prepared by", "Description"],
        [["V0.1", datetime.today().strftime("%Y-%m-%d"),
          _safe(data.get("author"), ""),
          "Initial Draft"]])

    doc.add_paragraph()

    # Document Information
    _bold_para(doc, "Document Information")
    _plain_kv_table(doc, [
        ("Document name:", f"{_safe(data.get('interface_id'), '')} - Technical Design"),
        ("Author:", _safe(data.get("author"), "")),
        ("Project Team:", _safe(data.get("project_team"), "")),
        ("Process Owner:", _safe(data.get("process_owner"), "")),
        ("E-mail:", ""),
        ("Telephone:", ""),
        ("Organisation:", _safe(data.get("organisation"), "")),
        ("Location:", ""),
    ])

    doc.add_paragraph()

    # Reference Documents
    _bold_para(doc, "Reference Documents")
    _data_table(doc,
        ["Document name", "Description Document", "Version", "Link to document"],
        [["", "", "", ""]])

    doc.add_page_break()


def _system_table(doc: Document, sys_type: str, rows: list):
    """
    Source or Target system table — spanning dark header with sys_type label,
    then label-value rows. Matches the reference tables exactly.
    """
    _spanning_header_table(doc, sys_type, rows)


def build_td(data: dict) -> bytes:
    """
    Build a TD .docx document matching the reference format exactly.
    Sections: Cover, 1 General Information, 2 Technical Document Concept,
              3 Technical Design, 4 Operational Support, 5 ABAP Build Details,
              6 Unit Test, 7 Issue List, 8 Appendix.
    """
    doc = Document()

    # Page layout: 1-inch margins on all sides
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1)
        sec.right_margin  = Inches(1)

    _td_footer(doc)

    # ── COVER ─────────────────────────────────────────────────────────────────
    _cover_page(doc, data)

    # ── TABLE OF CONTENTS ─────────────────────────────────────────────────────
    _add_toc(doc)

    # ── SECTION 1: General Information ───────────────────────────────────────
    _h(doc, "1  General Information", 1)

    # 1.1
    _h(doc, "1.1  General Information - Integration", 2)
    _spanning_header_table(doc, "Working Name of Interface: " + _safe(data.get("interface_name"), ""), [
        ("Middleware Solutions:", _safe(data.get("middleware"), "CPI")),
        ("Business Owner:", _safe(data.get("business_owner"), "")),
        ("Business criticality:", _safe(data.get("business_criticality"), "Critical")),
        ("Source system:", _safe(data.get("source_system"), "")),
        ("Target system:", _safe(data.get("target_system"), "")),
        ("Data flow direction:", _safe(data.get("data_flow"), "Unidirectional")),
        ("List of data objects that need to be transferred:",
         _join_list(data.get("data_objects", ""))),
    ])

    # 1.2
    _h(doc, "1.2  Conditions", 2)
    _plain_kv_table(doc, [
        ("Planned start date of testing:", _safe(data.get("planned_test_date"), "TBD")),
        ("Planned go-live date:", _safe(data.get("planned_golive_date"), "TBD")),
    ])
    _bold_para(doc, "Description of business process + EA diagram:")
    doc.add_paragraph(_safe(data.get("business_process_description"), ""))

    # 1.3
    _h(doc, "1.3  Assumptions", 2)
    assumptions = _ensure_list(data.get("assumptions", "N/A"))
    for a in (assumptions or ["N/A"]):
        _bullet(doc, a)

    # 1.4
    _h(doc, "1.4  Involved Systems, Message Communication and Interfaces", 2)

    # 1.4.1
    _h(doc, "1.4.1  Involved Systems", 3)
    _system_table(doc, "Source", [
        ("Name of application:", _safe(data.get("source_app_name"), "")),
        ("Use of System:", _safe(data.get("source_use"), "")),
        ("Type of System:", _safe(data.get("source_type"), "SAP System")),
        ("Communication Protocol:", _safe(data.get("source_protocol"), "")),
    ])
    _system_table(doc, "Target", [
        ("Name of application:", _safe(data.get("target_app_name"), "")),
        ("Use of System:", _safe(data.get("target_use"), "")),
        ("Type of System:", _safe(data.get("target_type"), "Third-Party System")),
        ("Communication Protocol:", _safe(data.get("target_protocol"), "")),
    ])

    # 1.4.2
    _h(doc, "1.4.2  Message Communication", 3)
    _spanning_header_table(doc, "Name of the message", [
        ("Content:", _safe(data.get("message_content"), "")),
        ("Sample File (source + target):", _safe(data.get("message_sample"), "N/A")),
        ("Frequency:", _safe(data.get("message_frequency"), "")),
        ("Estimated Size:", _safe(data.get("message_size"), "TBD")),
    ])

    # 1.4.3
    _h(doc, "1.4.3  Interfaces", 3)
    _spanning_header_table(doc, "Dependencies", [
        ("Are there any dependencies to other interfaces known?",
         "Yes" if data.get("has_dependencies") else "No"),
        ("If yes, specify — related interface (name/ID):",
         _safe(data.get("dependencies_detail"), "N/A")),
        ("Type of dependency:", "N/A"),
    ])
    _spanning_header_table(doc, "Security", [
        ("Does the message contain sensitive data?",
         "Yes" if data.get("sensitive_data") else "No"),
        ("Transport level encryption:", _safe(data.get("transport_encryption"), "TLS 1.2 or higher")),
        ("Message level encryption:", _safe(data.get("message_encryption"), "N/A")),
    ])
    _spanning_header_table(doc, "Authorisation", [
        ("What authorisation does the interface user require for executing the described "
         "activities in the receiving backend system?",
         _safe(data.get("auth_receiving"), "User: Placeholder  Profile: Placeholder")),
        ("What authorisation / role does the interface user require for connecting "
         "to the integration solution?",
         _safe(data.get("auth_integration"), "User: Placeholder  User role: Placeholder")),
    ])

    doc.add_page_break()

    # ── SECTION 2: Technical Document Concept ────────────────────────────────
    _h(doc, "2  Technical Document Concept", 1)

    _h(doc, "2.1  General Information - Complexity", 2)
    _plain_kv_table(doc, [
        ("Complexity of Development:", _safe(data.get("complexity"), "Medium")),
        ("Integration Developer:", _safe(data.get("developer"), "TBD")),
    ])

    _h(doc, "2.2  Interface – Technical Design", 2)
    _spanning_header_table(doc, "Integration logic", [
        ("Integration flow:", _safe(data.get("integration_logic"))),
        ("Flowchart of integration logic:", "See diagram below"),
    ])

    # ── Flow diagram ──────────────────────────────────────────────────────────
    try:
        png_bytes = generate_flowchart(data)
        doc.add_picture(BytesIO(png_bytes), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Flow diagram could not be generated: {e}]")

    doc.add_page_break()

    # ── SECTION 3: Technical Design ──────────────────────────────────────────
    _h(doc, "3  Technical Design", 1)
    _h(doc, "3.1  Interface Objects", 2)

    _h(doc, "3.1.1  Interface", 3)
    _spanning_header_table(doc, "IFlow (PO) or Artifact (CPI)", [
        ("Folder name (PO) or Package (CPI):", _safe(data.get("package_name"), "")),
        ("IFlow (PO) or Artifact (CPI):", _safe(data.get("artifact_name"), "")),
        ("Mode:", _safe(data.get("mode"), "Asynchronous")),
        ("Description:", _safe(data.get("iflow_description"), "")),
    ])

    _h(doc, "3.1.2  Mapping Objects", 3)
    mapping_objects = data.get("mapping_objects")
    if mapping_objects and isinstance(mapping_objects, list) and len(mapping_objects) > 0:
        rows = []
        for mo in mapping_objects:
            if isinstance(mo, (list, tuple)) and len(mo) >= 4:
                rows.append(list(mo[:4]))
            elif isinstance(mo, dict):
                rows.append([
                    mo.get("name", ""), mo.get("type", ""),
                    mo.get("namespace", ""), mo.get("specification", ""),
                ])
        _data_table(doc,
            ["Name", "Type", "Namespace / Artifact", "Mapping Specification"],
            rows)
    else:
        _spanning_header_table(doc, "Mapping", [
            ("Name:", _safe(data.get("mapping_name"), "")),
            ("Namespace (PO) or Artifact (CPI):", _safe(data.get("mapping_namespace"), "")),
            ("Software Component Version (PO only):", "N/A"),
            ("Mapping Specification:", _safe(data.get("mapping_spec"), "")),
            ("Mapping type (Message Mapping, Groovy, XSLT, JavaScript, JAVA):",
             _safe(data.get("mapping_type"), "Message Mapping")),
        ])

    doc.add_page_break()

    # ── SECTION 4: Operational Support ───────────────────────────────────────
    _h(doc, "4  Operational Support", 1)
    _h(doc, "4.1  Monitoring", 2)

    _spanning_header_table(doc, "Within Sender", [
        ("Monitoring Solution:", _safe(data.get("sender_monitoring"), "TBD")),
        ("Contact information:", _safe(data.get("sender_contact"), "The responsible team")),
        ("Additional information:", "TBD"),
    ])
    _spanning_header_table(doc, "Within PO / CPI", [
        ("Monitoring solution:", _safe(data.get("cpi_monitoring"),
            "Message failures will be sent to operation team")),
        ("Contact information:", _safe(data.get("cpi_contact"), "Developer / Operational Team (TBF)")),
        ("Additional information:", _safe(data.get("cpi_alert"),
            "Mapping Failures cannot be retriggered. If system is down, messages will be pushed again.")),
        ("Cloud ALM fields:", _safe(data.get("cloud_alm"),
            "Property Name - SAP_MessageProcessingLogCustomStatus : "
            "Value - COMPLETED ${header.SAP_ApplicationID}")),
    ])
    _spanning_header_table(doc, "Within Receiver", [
        ("Monitoring Solution:", _safe(data.get("receiver_monitoring"), "TBD")),
        ("Contact information:", _safe(data.get("receiver_contact"), "Third party contact")),
        ("Additional information:", "TBD"),
    ])

    _h(doc, "4.2  Unit Test / UAT", 2)
    _spanning_header_table(doc, "Unit Test", [
        ("Describe how a unit test can be performed:",
         _safe(data.get("unit_test_description"), "")),
        ("Any relevant materials, such as Postman collection, test files:",
         _safe(data.get("unit_test_materials"), "")),
    ])

    _h(doc, "4.3  RFC history", 2)
    _data_table(doc, ["Ticket", "Description"], [("N/A", "N/A")])

    _h(doc, "4.4  ISSUE LIST:", 2)
    issues = data.get("issues", [])
    if issues:
        rows = []
        for iss in issues:
            if isinstance(iss, (list, tuple)):
                rows.append([str(iss[0]) if len(iss) > 0 else "",
                             str(iss[1]) if len(iss) > 1 else ""])
            elif isinstance(iss, dict):
                rows.append([iss.get("issue", ""), iss.get("resolution", "")])
        _data_table(doc, ["Issue / Failure", "Resolution"], rows)
    else:
        _data_table(doc, ["Issue / Failure", "Resolution"],
                    [("No open issues.", "")])

    doc.add_page_break()

    # ── SECTION 5: ABAP Build Details ────────────────────────────────────────
    _h(doc, "5  ABAP Build details", 1)
    _h(doc, "5.1  ABAP Build", 2)

    _h(doc, "5.1.1  Transport Details", 3)
    _data_table(doc,
        ["Transport ID", "Transport Description", "Transport Created By"],
        [("N/A", "N/A", "N/A")])

    _h(doc, "5.1.2  Open Items", 3)
    _data_table(doc,
        ["Date Opened", "Description", "Status", "Assigned To", "Due Date"],
        [("N/A", "N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.1.3  Configuration and System Setup", 3)
    _data_table(doc,
        ["Configuration / System Setup", "Related T.Code", "Configured / Setup Values"],
        [("N/A", "N/A", "N/A")])

    _h(doc, "5.1.4  RICEFW Build Components", 3)
    _data_table(doc,
        ["Build Component/Object", "Type of Development", "Detail of the Development", "Status"],
        [("N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.1.5  Processing", 3)
    proc_list = _ensure_list(data.get("processing_description",
                                       data.get("processing", "N/A")))
    for p in (proc_list or ["N/A"]):
        _bullet(doc, p)

    _h(doc, "5.1.6  Processing Type", 3)
    proc_type_list = _ensure_list(data.get("processing_type_detail",
                                            data.get("processing_type", "N/A")))
    for p in (proc_type_list or ["N/A"]):
        _bullet(doc, p)

    _h(doc, "5.1.7  Enhanced Rules for Standard Segments/Fields", 3)
    enhanced = data.get("enhanced_fields", [])
    if enhanced and isinstance(enhanced, list) and len(enhanced) > 0:
        rows = []
        for ef in enhanced:
            if isinstance(ef, (list, tuple)):
                rows.append(list(ef[:5]) + [""] * max(0, 5 - len(ef)))
            elif isinstance(ef, dict):
                rows.append([
                    ef.get("segment", ""), ef.get("field", ""),
                    ef.get("description", ""), ef.get("qualifier", ""),
                    ef.get("logic", ""),
                ])
        _data_table(doc,
            ["Segment Name", "Field", "Field Description", "Qualifier", "Logic"],
            rows)
    else:
        _data_table(doc,
            ["Segment Name", "Field", "Field Description", "Qualifier", "Logic"],
            [("N/A", "N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.1.8  Custom Segment Names and Descriptions", 3)
    _data_table(doc,
        ["Parent Segment", "Custom Segment Name", "Custom Segment Description", "Segment Occurence"],
        [("N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.1.9  Custom Segment Fields", 3)
    _data_table(doc,
        ["Segment Name", "Field", "Field Description", "Qualifier", "Logic"],
        [("N/A", "N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.1.10  Cross Reference/Table Maintenance", 3)
    doc.add_paragraph("N/A")

    _h(doc, "5.1.11  Queue Management", 3)
    _bullet(doc, "N/A")

    _h(doc, "5.2  Preconditions / Considerations", 2)

    _h(doc, "5.2.1  Master Data Entries", 3)
    md_list = _ensure_list(data.get("master_data", "N/A"))
    for m in (md_list or ["N/A"]):
        _bullet(doc, m)

    _h(doc, "5.2.2  Transaction Data", 3)
    _bullet(doc, "N/A")

    _h(doc, "5.2.3  Customizing", 3)
    doc.add_paragraph("N/A")

    _h(doc, "5.2.4  Security Considerations", 3)
    sc_list = _ensure_list(data.get("security_considerations", "N/A"))
    for s in (sc_list or ["N/A"]):
        _bullet(doc, s)

    _h(doc, "5.2.5  Exception Handling", 3)
    eh_list = _ensure_list(data.get("error_handling", "N/A"))
    for e in (eh_list or ["N/A"]):
        _bullet(doc, e)

    _h(doc, "5.2.6  Message/Error Handling", 3)
    _data_table(doc,
        ["Code", "Type", "Action", "Description"],
        [("N/A", "N/A", "N/A", "N/A")])

    _h(doc, "5.3  Development Considerations", 2)
    dev_rows = data.get("dev_considerations", [])
    if dev_rows and isinstance(dev_rows, list) and len(dev_rows) > 0:
        rows = []
        for dr in dev_rows:
            if isinstance(dr, (list, tuple)) and len(dr) >= 2:
                rows.append([str(dr[0]), str(dr[1])])
            elif isinstance(dr, dict):
                rows.append([dr.get("area", ""), dr.get("details", "")])
        _data_table(doc, ["Development Considerations", "Details and Values"], rows)
    else:
        _data_table(doc, ["Development Considerations", "Details and Values"], [
            ("Integration",          _safe(data.get("dev_integration"), "N/A")),
            ("Frequency and Transaction Volume",
             _safe(data.get("dev_frequency"), "N/A")),
            ("User Roles and Authorization",
             "Only system integration user requires access"),
            ("Operation and Performance",
             _safe(data.get("dev_performance"), "N/A")),
            ("Security",             _safe(data.get("dev_security"), "TLS 1.2 or higher")),
            ("Execution and Usage",  _safe(data.get("dev_execution"), "N/A")),
        ])

    _h(doc, "5.3.1  Integration Component", 3)
    _data_table(doc,
        ["List of Integration Components", "Details of Integration component"], [
        ("Events",  "Integration creation/change events in source system"),
        ("Transformations",
         _safe(data.get("mapping_type"), "Message Mapping")),
        ("Adapters",
         f"{_safe(data.get('source_protocol'), 'N/A')} / {_safe(data.get('target_protocol'), 'N/A')}"),
        ("Queues",  "N/A"),
        ("Automation Business Process Model", "N/A"),
        ("Workflow Model", "N/A"),
        ("Business Activity Monitoring Views",
         "Integration monitoring and error notification handled in CPI"),
    ])

    _h(doc, "5.3.2  Peer and QA Review", 3)
    _data_table(doc,
        ["Reviewer Type", "Review Date", "Reviewer Name",
         "Coding Standard Reviewed", "Security Standard reviewed"], [
        ("Peer Review", "", "",
         "No", "No"),
        ("QA Review",   "", "",
         "No", "No"),
    ])

    doc.add_page_break()

    # ── SECTION 6: UNIT TEST ──────────────────────────────────────────────────
    _h(doc, "6  UNIT TEST", 1)

    _spanning_header_table(doc, "Unit Test", [
        ("Describe how a unit test can be performed:",
         _safe(data.get("unit_test_description"), "")),
        ("Any relevant materials, such as Postman collection, test files:",
         _safe(data.get("unit_test_materials"), "")),
    ])

    _h(doc, "6.1  Test Tools and Environment", 2)
    _plain_kv_table(doc, [
        ("Test System",    _safe(data.get("test_system"), "")),
        ("Client",         _safe(data.get("test_client"), "")),
        ("Tool",           _safe(data.get("test_tool"), "")),
        ("Test Program(s)",""),
    ])

    _h(doc, "6.2  Technical Test Conditions", 2)
    test_conds = _ensure_list(data.get("technical_test_conditions", []))
    for t in (test_conds or [""]):
        _bullet(doc, t)

    _h(doc, "6.3  Unit Test Scenario & Results", 2)
    unit_scenarios = _ensure_list(
        data.get("unit_test_scenarios", data.get("test_scenarios", [])))
    for s in (unit_scenarios or [""]):
        _bullet(doc, s)

    doc.add_page_break()

    # ── SECTION 7: Issue List ─────────────────────────────────────────────────
    _h(doc, "7  Issue List", 1)
    if issues:
        rows = []
        for iss in issues:
            if isinstance(iss, (list, tuple)):
                rows.append([str(iss[0]) if len(iss) > 0 else "",
                             str(iss[1]) if len(iss) > 1 else ""])
            elif isinstance(iss, dict):
                rows.append([iss.get("issue", ""), iss.get("resolution", "")])
        _data_table(doc, ["Issue / Failure", "Resolution"], rows)
    else:
        _data_table(doc, ["Issue / Failure", "Resolution"],
                    [("No open issues.", "")])

    doc.add_page_break()

    # ── SECTION 8: Appendix ───────────────────────────────────────────────────
    _h(doc, "8  Appendix", 1)
    _h(doc, "8.1  information", 2)

    _h(doc, "General Information – Integration", 3)
    _spanning_header_table(doc, "Working Name of Interface: " + _safe(data.get("interface_name"), ""), [
        ("Middleware Solutions (CPI/PO/API Management/ Event Mesh/ B2B Framework):",
         _safe(data.get("middleware"), "CPI")),
        ("Business Owner:", _safe(data.get("business_owner"), "")),
        ("Business criticality (Highly Critical/critical/Non-critical):",
         _safe(data.get("business_criticality"), "Critical")),
        ("Source system:", _safe(data.get("source_system"), "")),
        ("Target system:", _safe(data.get("target_system"), "")),
        ("Data flow direction (Unidirectional/Bidirectional/Multiple):",
         _safe(data.get("data_flow"), "Unidirectional")),
        ("List of data objects that need to be transferred:",
         _join_list(data.get("data_objects", ""))),
    ])

    _h(doc, "Table 2:", 3)
    _plain_kv_table(doc, [
        ("Planned start date of testing:", _safe(data.get("planned_test_date"), "TBD")),
        ("Planned go-live date:",          _safe(data.get("planned_golive_date"), "TBD")),
    ])
    _bold_para(doc, "Description of business process + EA diagram:")
    doc.add_paragraph(_safe(data.get("business_process_description"), ""))

    _h(doc, "Involved Systems:", 3)
    _h(doc, "Source Table:", 4)
    _system_table(doc, "Source", [
        ("Name of application:", _safe(data.get("source_app_name"), "")),
        ("Use of System:", _safe(data.get("source_use"), "")),
        ("Type of System: (SAP System / Non-SAP/B2B)",
         _safe(data.get("source_type"), "SAP System")),
        ("Communication Protocol (IDoc, RFC, XI, JDBC, AS2, AMQP, SFTP, HTTPS, SOAP, Other):",
         _safe(data.get("source_protocol"), "")),
    ])
    _h(doc, "Target Table", 4)
    _system_table(doc, "Target", [
        ("Name of application:", _safe(data.get("target_app_name"), "")),
        ("Use of System:", _safe(data.get("target_use"), "")),
        ("Type of System: (SAP System / Non-SAP/B2B)",
         _safe(data.get("target_type"), "")),
        ("Communication Protocol (IDoc, RFC, XI, JDBC, AS2, AMQP, SFTP, HTTPS, SOAP, Other):",
         _safe(data.get("target_protocol"), "")),
    ])

    _h(doc, "Interfaces", 3)
    _h(doc, "Dependencies", 4)
    _spanning_header_table(doc, "Dependencies", [
        ("Are there any dependencies to other interfaces known?",
         "Yes" if data.get("has_dependencies") else "No"),
        ("Specify related interface (name/ID) if there are any dependency:",
         _safe(data.get("dependencies_detail"), "N/A")),
        ("Type of dependency:", "N/A"),
    ])
    _h(doc, "Security", 4)
    _spanning_header_table(doc, "Security", [
        ("Does the message contain sensitive data?",
         "Yes" if data.get("sensitive_data") else "No"),
        ("Transport level encryption:", _safe(data.get("transport_encryption"), "TLS 1.2 or higher")),
        ("Message level encryption:", _safe(data.get("message_encryption"), "N/A")),
    ])
    _h(doc, "Authorisation", 4)
    _spanning_header_table(doc, "Authorisation", [
        ("What authorisation does the interface user require for executing the described "
         "activities in the receiving backend system?",
         _safe(data.get("auth_receiving"), "User: Placeholder  Profile: Placeholder")),
        ("What authorisation / role does the interface user require for connecting "
         "to the integration solution?",
         _safe(data.get("auth_integration"), "User: Placeholder  User role: Placeholder")),
    ])

    _h(doc, "Complexity", 3)
    _plain_kv_table(doc, [
        ("Complexity of Development:", _safe(data.get("complexity"), "Medium")),
        ("Integration Developer:", _safe(data.get("developer"), "TBD")),
        ("e-mail:", ""),
    ])

    _h(doc, "IFlow(PO) or Artifact(CPI)", 3)
    _spanning_header_table(doc, "IFlow (PO) or Artifact (CPI)", [
        ("Folder name (PO) or Package (CPI):", _safe(data.get("package_name"), "")),
        ("IFlow (PO) or Artifact (CPI):", _safe(data.get("artifact_name"), "")),
        ("Mode (Asynchronous, Asynchronous via Event Mesh, Synchronous):",
         _safe(data.get("mode"), "Asynchronous")),
        ("Description:", _safe(data.get("iflow_description"), "")),
    ])

    _h(doc, "Mapping Objects", 3)
    _spanning_header_table(doc, "Mapping", [
        ("Name:", _safe(data.get("mapping_name"), "")),
        ("Namespace (PO) or Artifact (CPI):", _safe(data.get("mapping_namespace"), "")),
        ("Software Component Version (PO only):", "N/A"),
        ("Step 1 Type (Message Mapping, Groovy, XSLT, JavaScript, JAVA):",
         _safe(data.get("mapping_type"), "Message Mapping")),
        ("Name:", _safe(data.get("mapping_name"), "")),
        ("Additional Information:", "Mapping specification per FD"),
        ("Mapping Specification:", _safe(data.get("mapping_spec"), "")),
    ])

    _h(doc, "Sender Table", 3)
    _spanning_header_table(doc, "Within Sender", [
        ("Monitoring Solution:", _safe(data.get("sender_monitoring"), "TBD")),
        ("Contact information:", _safe(data.get("sender_contact"), "The responsible team")),
        ("Additional information:", "TBD"),
    ])

    _h(doc, "Within PO/CPI Table", 3)
    _spanning_header_table(doc, "Within PO / CPI", [
        ("Monitoring solution:", _safe(data.get("cpi_monitoring"), "TBD")),
        ("Contact information:", _safe(data.get("cpi_contact"), "Developer / Operational Team (TBF)")),
        ("Additional information:", _safe(data.get("cpi_alert"),
            "Mapping Failures cannot be retriggered.")),
        ("Cloud ALM fields:", _safe(data.get("cloud_alm"),
            "Property Name - SAP_MessageProcessingLogCustomStatus : "
            "Value - COMPLETED ${header.SAP_ApplicationID}")),
    ])

    _h(doc, "Receiver Table", 3)
    _spanning_header_table(doc, "Within Receiver", [
        ("Monitoring Solution:", _safe(data.get("receiver_monitoring"), "TBD")),
        ("Contact information:", _safe(data.get("receiver_contact"), "Third party contact")),
        ("Additional information:", "TBD"),
    ])

    _h(doc, "Unit Test", 3)
    _spanning_header_table(doc, "Unit Test", [
        ("Describe how a unit test can be performed:",
         _safe(data.get("unit_test_description"), "")),
        ("Any relevant materials, such as Postman collection, test files:",
         _safe(data.get("unit_test_materials"), "")),
    ])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
