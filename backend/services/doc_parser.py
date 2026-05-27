from docx import Document
from io import BytesIO
from typing import TypedDict

class DocImage(TypedDict):
    bytes: bytes
    content_type: str
    size: int


def parse_docx_to_text(file_bytes: bytes) -> dict:
    """Parse a .docx file and return text grouped by heading sections."""
    doc = Document(BytesIO(file_bytes))
    sections = {}
    current_heading = "header"
    current_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading" in style:
            if current_text:
                sections[current_heading] = "\n".join(current_text)
            current_heading = text
            current_text = []
        else:
            current_text.append(text)

    if current_text:
        sections[current_heading] = "\n".join(current_text)

    # Also extract all table data
    tables_text = []
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            seen = []
            for cell in row.cells:
                ct = cell.text.strip().replace("\n", " ")
                if ct and ct not in seen:
                    seen.append(ct)
            if any(seen):
                rows.append(" | ".join(seen))
        if rows:
            tables_text.append(f"Table {i+1}:\n" + "\n".join(rows))

    sections["_tables"] = "\n\n".join(tables_text)
    return sections


def extract_images_from_docx(file_bytes: bytes, min_size: int = 8000) -> list[DocImage]:
    """
    Extract embedded images from a .docx file.
    min_size: skip tiny images (icons, logos) below this byte threshold.
    Returns list sorted largest-first (most likely to be a flow diagram).
    """
    doc = Document(BytesIO(file_bytes))
    images: list[DocImage] = []

    for rel in doc.part.rels.values():
        if "image" not in rel.reltype.lower():
            continue
        try:
            part = rel.target_part
            blob: bytes = part.blob
            ct: str = part.content_type  # e.g. "image/png"
            if len(blob) >= min_size:
                images.append({"bytes": blob, "content_type": ct, "size": len(blob)})
        except Exception:
            pass

    # Also check images inside tables / inline shapes
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for rel in run._r.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        ):
                            rId = rel.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if rId and rId in doc.part.rels:
                                part = doc.part.rels[rId].target_part
                                blob = part.blob
                                if len(blob) >= min_size:
                                    entry = {"bytes": blob,
                                             "content_type": part.content_type,
                                             "size": len(blob)}
                                    if entry not in images:
                                        images.append(entry)

    return sorted(images, key=lambda x: x["size"], reverse=True)


def sections_to_summary(sections: dict) -> str:
    """Convert parsed sections to a readable summary string."""
    lines = []
    for heading, content in sections.items():
        if heading == "_tables":
            continue
        lines.append(f"## {heading}")
        lines.append(content)
        lines.append("")
    lines.append("## Tables Data")
    lines.append(sections.get("_tables", ""))
    return "\n".join(lines)
