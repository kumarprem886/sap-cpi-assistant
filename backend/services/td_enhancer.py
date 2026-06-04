"""Append a Developer Implementation Guide section to an existing TD document."""
import io
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from services.iflow_parser import parse_iflow_zip
from services.flowchart_builder import generate_flowchart

SAP_BLUE  = RGBColor(0x00, 0x6D, 0xB3)
SAP_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
AMBER     = RGBColor(0xCC, 0x77, 0x00)
RED_C     = RGBColor(0xC0, 0x2C, 0x2C)
GREEN     = RGBColor(0x1A, 0x8A, 0x5A)
MID_GREY  = RGBColor(0x55, 0x60, 0x70)

# Map activityType → human-readable label
STEP_LABELS = {
    'Script':                      'Groovy Script',
    'GroovyScript':                'Groovy Script',
    'Enricher':                    'Content Modifier',
    'ExternalCall':                'Request Reply',
    'Mapping':                     'Message Mapping',
    'ExclusiveGateway':            'Router (CBR)',
    'Splitter':                    'Splitter',
    'Gather':                      'Gather',
    'DBstorage':                   'DataStore',
    'StartTimerEvent':             'Timer Start',
    'ErrorEventSubProcessTemplate':'Exception Subprocess',
    'Send':                        'Send Step',
    '':                            'Step',
}

# Map activityType → palette location
PALETTE_LOCATION = {
    'Script':        'Call → Script → Groovy Script',
    'Enricher':      'Call → Content Modifier',
    'ExternalCall':  'Call → Request Reply',
    'Mapping':       'Call → Message Mapping',
    'ExclusiveGateway': 'Routing → Router',
    'Splitter':      'Call → Splitter → General Splitter',
    'DBstorage':     'Persistence → Write Variables → Data Store',
    'Send':          'Call → Send (for SFTP)',
    'StartTimerEvent': 'Events → Timer Start',
    'ErrorEventSubProcessTemplate': 'Call → Exception Subprocess',
}

def _shd(cell, hex_color):
    tc = cell._tc; p = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_color); s.set(qn('w:color'), 'auto'); s.set(qn('w:val'), 'clear')
    p.append(s)

def _tbl(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        r = c.paragraphs[0].runs[0]; r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        _shd(c, '1A1A2E')
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val) if val else ''
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i % 2 == 0: _shd(c, 'F2F4F7')
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows: row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return t

def _para(doc, text, bold=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text); run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color
    return p

def _code(doc, text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line or ' '); run.font.name = 'Courier New'; run.font.size = Pt(8); run.font.color.rgb = SAP_DARK

def _step_heading(doc, step_num, name, step_type, color='006DB3'):
    t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
    vals = [f'Step {step_num}', name, step_type]
    widths = [2, 10, 5]
    for j, (v, w) in enumerate(zip(vals, widths)):
        c = t.rows[0].cells[j]; c.text = v; c.width = Cm(w)
        r = c.paragraphs[0].runs[0]; r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)
        _shd(c, color)

def enhance_td_with_iflow(td_bytes: bytes, iflow_zip_bytes: bytes) -> bytes:
    """
    Load existing TD docx, parse iFlow ZIP, append Developer Guide section.
    Returns new docx bytes. Does NOT modify existing content.
    """
    # Load existing document
    doc = Document(io.BytesIO(td_bytes))

    # Parse iFlow
    data = parse_iflow_zip(iflow_zip_bytes)
    if not data:
        return td_bytes  # Return unchanged if parse fails

    # ── Page break before new section ────────────────────────────────────────
    doc.add_page_break()

    # ── Section header ────────────────────────────────────────────────────────
    h = doc.add_heading('Developer Implementation Guide', level=1)
    for run in h.runs:
        run.font.color.rgb = SAP_BLUE

    _para(doc,
        f'This section provides step-by-step configuration for recreating iFlow "{data["name"]}" '
        f'from scratch in SAP Integration Suite Designer. All palette steps, adapter properties, '
        f'Groovy scripts, and parameters are documented below.',
        size=10)

    # ── Flow diagram (SAP-themed PNG) ─────────────────────────────────────────
    doc.add_heading('Integration Flow Diagram', level=2)

    senders   = [a for a in data['adapters'] if a['direction'] == 'Sender']
    receivers = [a for a in data['adapters'] if a['direction'] == 'Receiver']

    # Build data dict for generate_flowchart
    src_name  = senders[0]['source_name']   if senders   else 'Source System'
    src_comp  = senders[0]['component']     if senders   else 'HTTPS'
    # Append adapter component to receiver name so the icon classifier can detect type
    # e.g. "Manogna" + "SFTP" → "Manogna SFTP" → shows folder icon
    tgt_names = ', '.join(
        f"{r['target_name']} {r['component']}" if r['component'] else r['target_name']
        for r in receivers
    ) if receivers else 'Target System'
    tgt_comp  = receivers[0]['component']   if receivers else 'HTTP'

    main_steps = [s for s in data['steps']
                  if s['element_type'] not in ('endEvent',)
                  and 'Error' not in s.get('id', '')
                  and 'Error' not in s.get('name', '')]

    step_desc = ', '.join(
        f"{s['name']} [{STEP_LABELS.get(s['activity_type'], s['element_type'])}]"
        for s in main_steps
    )
    activity_types = [s['activity_type'] for s in main_steps]
    mapping_type = (
        'Groovy' if any(a in ('Script', 'GroovyScript') for a in activity_types)
        else 'Message Mapping' if 'Mapping' in activity_types
        else ''
    )

    chart_data = {
        'interface_name':  data['name'],
        'source_app_name': src_name,
        'source_protocol': src_comp,
        'target_app_name': tgt_names,
        'target_protocol': tgt_comp,
        'integration_logic': step_desc,
        'mapping_type':    mapping_type,
    }

    try:
        png_bytes = generate_flowchart(chart_data)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
    except Exception as e:
        _para(doc, f'[Diagram could not be generated: {e}]', size=9, color=RED_C)

    # ── Adapters section ──────────────────────────────────────────────────────
    doc.add_heading('Adapters Configuration', level=2)

    for adapter in data['adapters']:
        direction = adapter['direction']
        name = adapter['target_name'] if direction == 'Receiver' else adapter['source_name']
        comp = adapter['component']
        color = '006DB3' if direction == 'Sender' else '1A1A2E'

        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        c0 = t.rows[0].cells[0]; c0.text = f"{direction}: {name}"; _shd(c0, color)
        c0.paragraphs[0].runs[0].font.bold = True; c0.paragraphs[0].runs[0].font.color.rgb = WHITE
        c0.paragraphs[0].runs[0].font.size = Pt(10); c0.width = Cm(8)
        c1 = t.rows[0].cells[1]; c1.text = f"Adapter: {comp}"; _shd(c1, color)
        c1.paragraphs[0].runs[0].font.bold = True; c1.paragraphs[0].runs[0].font.color.rgb = WHITE
        c1.paragraphs[0].runs[0].font.size = Pt(10); c1.width = Cm(9)

        # Key properties table
        props = adapter['properties']
        key_props = [(k, v) for k, v in props.items()
                     if v and k not in ('ComponentNS','ComponentSWCVName','ComponentSWCVId',
                                         'TransportProtocol','MessageProtocol','ComponentType',
                                         'system','direction','TransportProtocolVersion',
                                         'MessageProtocolVersion','Name','Description')]
        if key_props:
            _tbl(doc, ['Property', 'Value'],
                 [(k, v if len(v) < 80 else v[:77]+'...') for k, v in key_props[:20]],
                 col_widths=[7, 10])
        else:
            doc.add_paragraph()

    # ── Steps section ─────────────────────────────────────────────────────────
    doc.add_heading('Step-by-Step Palette Configuration', level=2)
    _para(doc, 'Follow these steps in order in the SAP Integration Suite Designer.', size=10, color=MID_GREY)
    doc.add_paragraph()

    step_num = 1
    for step in data['steps']:
        elem_type     = step['element_type']
        activity_type = step['activity_type']
        name          = step['name']
        props         = step['properties']

        # Skip error handler inner elements shown elsewhere, keep subProcesses
        if elem_type == 'startEvent' and 'Error' in step['id']:
            continue
        if elem_type == 'endEvent' and 'Error' in step['id']:
            continue

        label    = STEP_LABELS.get(activity_type, elem_type.replace('Event','').title())
        palette  = PALETTE_LOCATION.get(activity_type, '')

        # Determine color
        if elem_type == 'startEvent':  color = '1A8A5A'
        elif elem_type == 'endEvent':  color = '1A8A5A'
        elif 'subprocess' in elem_type.lower() or 'Error' in label: color = 'C02C2C'
        elif elem_type == 'exclusiveGateway': color = 'CC7700'
        else: color = '1A1A2E'

        _step_heading(doc, step_num, name, label, color)
        step_num += 1

        if palette:
            _para(doc, f'Palette: {palette}', size=9, color=MID_GREY, indent=0.3)

        # Show important properties
        show_props = {k: v for k, v in props.items()
                      if v and k not in ('cmdVariantUri','componentVersion','activityType',
                                          'subActivityType','scriptBundleId')}

        if show_props:
            _tbl(doc, ['Property', 'Value'],
                 [(k, v if len(v) < 90 else v[:87]+'...') for k, v in list(show_props.items())[:15]],
                 col_widths=[6, 11])

        # Groovy script code
        if step['script_file'] and step['script_code']:
            _para(doc, f'Script file: {step["script_file"]}  |  Function: {step["script_fn"]}',
                  size=9, color=MID_GREY, indent=0.3)
            _para(doc, 'Groovy Script Code:', bold=True, size=10)
            _code(doc, step['script_code'][:3000])  # cap at 3000 chars

        # Exception subprocess inner steps
        if step['inner_steps']:
            _para(doc, 'Inner steps:', bold=True, size=10)
            for inner in step['inner_steps']:
                _para(doc, f"  - {inner['name']} ({STEP_LABELS.get(inner['activity_type'], inner['element_type'])})",
                      size=9)

    # ── Parameters section ────────────────────────────────────────────────────
    if data['parameters']:
        doc.add_heading('Externalized Parameters', level=2)
        _para(doc, 'After import: Configure → Parameters tab. Fill in real values before deploying.',
              size=10, color=MID_GREY)
        _tbl(doc, ['Parameter Name', 'Default Value', 'Required', 'Description'],
             [(p['name'], p['default'], 'Yes' if p['required'] else 'No', p['description'])
              for p in data['parameters']],
             col_widths=[5, 5, 2.5, 5])

    # ── How to deploy ─────────────────────────────────────────────────────────
    doc.add_heading('Deployment Steps', level=2)
    steps_list = [
        'Import the iFlow ZIP into your integration package (CPI Connect → Import ZIP)',
        'Go to Configure → Parameters tab → fill in all real values',
        'Create security artifacts (credentials/keys) referenced by adapter configs',
        'Deploy to SANDBOX first for testing',
        'Run test messages → check Message Monitor for COMPLETED status',
        'Fix any errors, re-deploy → test again',
        'Deploy to PRODUCTION only after sandbox testing passes',
    ]
    for i, s in enumerate(steps_list, 1):
        _para(doc, f'{i}. {s}', size=10)

    # Save to bytes
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# STANDALONE TD FROM IFLOW — ZERO AI
# Builds a complete TD document from an iFlow ZIP alone.
# Technical sections (steps, adapters, scripts, params) are 100% accurate
# because they come directly from the parsed iFlow XML.
# Business sections (description, assumptions, test plan) are TBD placeholders
# for the developer to fill in.
# ─────────────────────────────────────────────────────────────────────────────

def build_td_from_iflow(iflow_zip_bytes: bytes, author: str = "", project_team: str = "") -> bytes:
    """
    Generate a complete Technical Design document from an iFlow ZIP.
    No AI involved — 100% deterministic from the iFlow XML.

    Sections produced:
      Cover + metadata
      1. Integration Overview  (from participants + adapters)
      2. Flow Diagram          (SAP-themed PNG)
      3. Adapters              (all connection properties)
      4. Step-by-Step Config   (every palette step)
      5. Groovy Scripts        (full code)
      6. Parameters            (all externalized params with defaults)
      7. Security Prerequisites
      8. Test Plan             (placeholder)
    """
    data = parse_iflow_zip(iflow_zip_bytes)
    if not data:
        raise ValueError("Could not parse iFlow ZIP")

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2.0)
        sec.left_margin = sec.right_margin = Cm(2.0)

    senders   = [a for a in data['adapters'] if a['direction'] == 'Sender']
    receivers = [a for a in data['adapters'] if a['direction'] == 'Receiver']
    scripts   = data.get('scripts', {})
    params    = data.get('parameters', [])

    # ── Cover page ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run('TECHNICAL DESIGN DOCUMENT'); r.font.size = Pt(24); r.font.bold = True; r.font.color.rgb = SAP_BLUE

    p2 = doc.add_paragraph(); r2 = p2.add_run(data['name'])
    r2.font.size = Pt(18); r2.font.bold = True; r2.font.color.rgb = SAP_DARK

    _para(doc, data.get('description', ''), size=12, color=MID_GREY)
    doc.add_paragraph()

    _tbl(doc, ['Field', 'Value'], [
        ('iFlow Name',    data['name']),
        ('Author',        author or 'TBD'),
        ('Project Team',  project_team or 'TBD'),
        ('Middleware',    'SAP Integration Suite (CPI)'),
        ('Complexity',    _detect_complexity(data)),
        ('Generated',     'Automatically from iFlow XML — no AI'),
        ('Status',        'Draft — fill in business sections'),
    ], col_widths=[5, 12])

    doc.add_page_break()

    # ── Section 1: Integration Overview ───────────────────────────────────────
    h = doc.add_heading('1.  Integration Overview', level=1)
    for r in h.runs: r.font.color.rgb = SAP_BLUE

    _para(doc, data.get('description') or 'Fill in the business description of this interface.', size=10)
    doc.add_paragraph()

    # Sender / receiver overview table
    src_name = senders[0]['source_name']   if senders   else 'Source'
    src_comp = senders[0]['component']     if senders   else '—'
    _tbl(doc, ['Attribute', 'Value'], [
        ('Source System',     src_name),
        ('Sender Adapter',    src_comp),
        ('Target System(s)',  ', '.join(r['target_name'] for r in receivers) if receivers else 'TBD'),
        ('Receiver Adapters', ', '.join(sorted({r['component'] for r in receivers})) if receivers else '—'),
        ('Mode',              'Asynchronous'),
        ('Direction',         'Unidirectional' if len(receivers) <= 1 else 'Fan-out (1 → N)'),
    ], col_widths=[5, 12])

    # ── Section 2: Flow Diagram ────────────────────────────────────────────────
    h2 = doc.add_heading('2.  Integration Flow Diagram', level=1)
    for r in h2.runs: r.font.color.rgb = SAP_BLUE

    _para(doc, 'Diagram generated deterministically from iFlow XML — accurate to the actual configuration.', size=9, color=MID_GREY)

    tgt_names_for_chart = ', '.join(
        f"{r['target_name']} {r['component']}" for r in receivers
    ) if receivers else 'Target'

    main_steps = [s for s in data['steps']
                  if s['element_type'] not in ('endEvent',)
                  and 'Error' not in s.get('id','')
                  and 'Error' not in s.get('name','')]
    step_desc = ', '.join(
        f"{s['name']} [{STEP_LABELS.get(s['activity_type'], s['element_type'])}]"
        for s in main_steps
    )
    activity_types = [s['activity_type'] for s in main_steps]
    mapping_type = (
        'Groovy' if any(a in ('Script','GroovyScript') for a in activity_types)
        else 'Message Mapping' if 'Mapping' in activity_types else ''
    )

    try:
        png_bytes = generate_flowchart({
            'interface_name':  data['name'],
            'source_app_name': src_name,
            'source_protocol': src_comp,
            'target_app_name': tgt_names_for_chart,
            'target_protocol': receivers[0]['component'] if receivers else '',
            'integration_logic': step_desc,
            'mapping_type':    mapping_type,
        })
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(png_bytes), width=Inches(6.2))
    except Exception as e:
        _para(doc, f'[Diagram error: {e}]', size=9, color=RED_C)

    doc.add_page_break()

    # ── Section 3: Adapters ────────────────────────────────────────────────────
    h3 = doc.add_heading('3.  Adapter Configuration', level=1)
    for r in h3.runs: r.font.color.rgb = SAP_BLUE

    for adapter in data['adapters']:
        direction = adapter['direction']
        sys_name  = adapter['source_name'] if direction == 'Sender' else adapter['target_name']
        comp      = adapter['component']
        color     = '006DB3' if direction == 'Sender' else '1A1A2E'

        t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
        for j, (txt, w) in enumerate([(direction, 2.5),(f"{sys_name}", 8),(comp, 7)]):
            c = t.rows[0].cells[j]; c.text = txt; c.width = Cm(w)
            r = c.paragraphs[0].runs[0]; r.font.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(10)
            _shd(c, color)

        props = adapter['properties']
        show  = {k: v for k, v in props.items()
                 if v and k not in ('ComponentNS','ComponentSWCVName','ComponentSWCVId',
                                     'system','direction','ComponentType','Description')}
        if show:
            _tbl(doc, ['Property', 'Value'],
                 [(k, v[:90] if len(v)>90 else v) for k,v in list(show.items())[:25]],
                 col_widths=[6, 11.5])

    # ── Section 4: Step-by-Step Palette Configuration ─────────────────────────
    h4 = doc.add_heading('4.  Step-by-Step Palette Configuration', level=1)
    for r in h4.runs: r.font.color.rgb = SAP_BLUE
    _para(doc, 'Follow these steps in order in SAP Integration Suite Designer to recreate this iFlow.', size=10, color=MID_GREY)
    doc.add_paragraph()

    step_num = 1
    for step in data['steps']:
        elem      = step['element_type']
        act       = step['activity_type']
        name      = step['name']
        props     = step['properties']

        if elem == 'startEvent' and 'Error' in step['id']:  continue
        if elem == 'endEvent'   and 'Error' in step['id']:  continue

        label   = STEP_LABELS.get(act, elem.replace('Event','').title())
        palette = PALETTE_LOCATION.get(act, '')
        color   = ('1A8A5A' if 'Event' in elem
                   else 'C02C2C' if ('subprocess' in elem.lower() or 'Error' in label)
                   else 'CC7700' if elem == 'exclusiveGateway'
                   else '1A1A2E')

        _step_heading(doc, step_num, name, label, color)
        step_num += 1

        if palette:
            _para(doc, f'Palette location: {palette}', size=9, color=MID_GREY, indent=0.3)

        show_props = {k: v for k, v in props.items()
                      if v and k not in ('cmdVariantUri','componentVersion','activityType',
                                          'subActivityType','scriptBundleId')}
        if show_props:
            _tbl(doc, ['Property', 'Value'],
                 [(k, v[:90] if len(v)>90 else v) for k, v in list(show_props.items())[:15]],
                 col_widths=[6, 11.5])

        if step.get('script_file') and step.get('script_code'):
            _para(doc, f"Script: {step['script_file']}  |  Function: {step.get('script_fn','')}", size=9, color=MID_GREY, indent=0.3)

        if step.get('inner_steps'):
            for inner in step['inner_steps']:
                inner_label = STEP_LABELS.get(inner.get('activity_type',''), inner.get('element_type',''))
                _para(doc, f"  • {inner['name']} ({inner_label})", size=9)

    # ── Section 5: Groovy Scripts ──────────────────────────────────────────────
    if scripts:
        h5 = doc.add_heading('5.  Groovy Scripts', level=1)
        for r in h5.runs: r.font.color.rgb = SAP_BLUE

        for fname, code in scripts.items():
            t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
            c = t.rows[0].cells[0]; c.text = fname
            _shd(c, '1A1A2E')
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.color.rgb = WHITE
            c.paragraphs[0].runs[0].font.size = Pt(10)
            _code(doc, code[:4000])
            doc.add_paragraph()

    # ── Section 6: Parameters ──────────────────────────────────────────────────
    h6_num = 6 if scripts else 5
    h6 = doc.add_heading(f'{h6_num}.  Externalized Parameters', level=1)
    for r in h6.runs: r.font.color.rgb = SAP_BLUE
    _para(doc, 'After importing the iFlow: Configure → Parameters tab. Fill in all real values before deploying.', size=10, color=MID_GREY)

    if params:
        _tbl(doc, ['Parameter Name', 'Default Value', 'Required', 'Description'],
             [(p['name'], p['default'], 'Yes' if p['required'] else 'No', p['description'])
              for p in params],
             col_widths=[5, 5, 2.5, 5])
    else:
        _para(doc, 'No externalized parameters found.', size=10, color=MID_GREY)

    # ── Section 7: Security Prerequisites ─────────────────────────────────────
    h7_num = h6_num + 1
    h7 = doc.add_heading(f'{h7_num}.  Security Prerequisites', level=1)
    for r in h7.runs: r.font.color.rgb = SAP_BLUE

    # Extract credential aliases from adapter properties
    cred_aliases = []
    for adapter in data['adapters']:
        for k, v in adapter['properties'].items():
            if 'credential' in k.lower() and v and not v.startswith('{{'):
                cred_aliases.append((adapter.get('target_name') or adapter.get('source_name',''), k, v))
            elif v and v.startswith('{{') and 'Cred' in v:
                cred_aliases.append((adapter.get('target_name') or adapter.get('source_name',''), k, v))

    if cred_aliases:
        _tbl(doc, ['System', 'Property', 'Alias / Reference'],
             cred_aliases[:20], col_widths=[5, 5, 7.5])
        _para(doc, 'Create these security artifacts in CPI Security Material before deploying.', size=10, color=AMBER)
    else:
        _para(doc, 'Review adapter properties and create required security artifacts (credentials, keys).', size=10, color=MID_GREY)

    # ── Section 8: Deployment + Test Plan ─────────────────────────────────────
    h8_num = h7_num + 1
    h8 = doc.add_heading(f'{h8_num}.  Deployment & Test Plan', level=1)
    for r in h8.runs: r.font.color.rgb = SAP_BLUE

    _tbl(doc, ['Step', 'Action', 'Status'], [
        ('1', 'Import iFlow ZIP into integration package',          'TODO'),
        ('2', 'Configure all Parameters (URLs, credentials)',        'TODO'),
        ('3', 'Create security artifacts (credentials, key pairs)', 'TODO'),
        ('4', 'Deploy to SANDBOX — check startup status',           'TODO'),
        ('5', 'Send test message — check Message Monitor',          'TODO'),
        ('6', 'Validate output / target system received data',       'TODO'),
        ('7', 'Fix errors, re-test until COMPLETED with correct data','TODO'),
        ('8', 'Deploy to PRODUCTION',                               'TODO'),
    ], col_widths=[1, 12, 4.5])

    _para(doc, '\nTest scenarios — fill in based on business requirements:', bold=True, size=10)
    _para(doc, '• Happy path: normal message with valid data → expected target output', size=10, indent=0.5)
    _para(doc, '• Error path: invalid/missing mandatory field → error handled by Exception Subprocess', size=10, indent=0.5)
    _para(doc, '• Boundary test: empty message body, large payload, special characters', size=10, indent=0.5)

    # ── Save ──────────────────────────────────────────────────────────────────
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _detect_complexity(data: dict) -> str:
    """Estimate iFlow complexity from number of steps, scripts, adapters."""
    n_steps   = len(data.get('steps', []))
    n_scripts = len(data.get('scripts', {}))
    n_adapters = len(data.get('adapters', []))
    score = n_steps + n_scripts * 3 + n_adapters
    if score > 15: return 'High'
    if score > 7:  return 'Medium'
    return 'Low'
